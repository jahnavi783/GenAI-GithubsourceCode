# import os
# import sys
 
# # Add project root directory to Python path FIRST
# sys.path.append(os.path.dirname(os.path.dirname(__file__)))
# import env_loader  # noqa: F401
# import re
# import sys
# import threading
# import time
# import json
# from datetime import datetime
# from pathlib import Path

# import requests
# import streamlit as st

# try:
#     from streamlit_autorefresh import st_autorefresh
# except (ImportError, ModuleNotFoundError):
#     st_autorefresh = None

# # ── Page config ───────────────────────────────────────────────────────────────
# st.set_page_config(
#     page_title="DocSync",
#     page_icon="📄",
#     layout="wide",
#     initial_sidebar_state="expanded",
# )

# # ── Custom CSS ────────────────────────────────────────────────────────────────
# st.markdown("""
# <style>
# @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

# /* Global */
# html, body, [class*="css"] {
#     font-family: 'Inter', sans-serif;
# }

# /* Hide default streamlit elements */
# #MainMenu {visibility: hidden;}
# footer {visibility: hidden;}
# header {visibility: hidden;}
# .stDeployButton {display: none;}

# /* Main background */
# .stApp {
#     background: #f6f8fa;
#     color: #1f2328;
# }

# /* Sidebar */
# [data-testid="stSidebar"] {
#     background: #ffffff !important;
#     border-right: 1px solid #d0d7de !important;
# }
# [data-testid="stSidebar"] * {
#     color: #1f2328 !important;
# }

# /* Buttons */
# .stButton > button {
#     background: #1f883d !important;
#     color: #ffffff !important;
#     border: 1px solid #1a7f37 !important;
#     border-radius: 6px !important;
#     font-family: 'Inter', sans-serif !important;
#     font-weight: 500 !important;
#     padding: 8px 16px !important;
#     transition: all 0.2s !important;
# }
# .stButton > button:hover {
#     background: #1a7f37 !important;
#     border-color: #157832 !important;
# }

# /* Input fields */
# .stTextInput > div > div > input {
#     background: #ffffff !important;
#     border: 1px solid #d0d7de !important;
#     color: #1f2328 !important;
#     border-radius: 6px !important;
#     font-family: 'Inter', sans-serif !important;
# }
# .stTextInput > div > div > input:focus {
#     border-color: #0969da !important;
#     box-shadow: 0 0 0 3px rgba(9,105,218,0.12) !important;
# }

# /* Text area */
# .stTextArea textarea {
#     background: #ffffff !important;
#     border: 1px solid #d0d7de !important;
#     color: #1f2328 !important;
#     border-radius: 6px !important;
#     font-family: 'JetBrains Mono', monospace !important;
#     font-size: 13px !important;
# }

# /* Tabs */
# .stTabs [data-baseweb="tab-list"] {
#     background: transparent !important;
#     border-bottom: 1px solid #d0d7de !important;
#     gap: 0 !important;
# }
# .stTabs [data-baseweb="tab"] {
#     background: transparent !important;
#     color: #636c76 !important;
#     border: none !important;
#     padding: 8px 16px !important;
#     font-family: 'Inter', sans-serif !important;
#     font-size: 14px !important;
#     font-weight: 500 !important;
# }
# .stTabs [aria-selected="true"] {
#     color: #1f2328 !important;
#     border-bottom: 2px solid #fd8c73 !important;
# }

# /* Metric cards */
# [data-testid="stMetric"] {
#     background: #ffffff !important;
#     border: 1px solid #d0d7de !important;
#     border-radius: 8px !important;
#     padding: 16px !important;
# }
# [data-testid="stMetricLabel"] {
#     color: #636c76 !important;
#     font-size: 12px !important;
#     font-weight: 500 !important;
#     text-transform: uppercase !important;
#     letter-spacing: 0.05em !important;
# }
# [data-testid="stMetricValue"] {
#     color: #1f2328 !important;
#     font-size: 24px !important;
#     font-weight: 700 !important;
# }

# /* Divider */
# hr {
#     border-color: #d0d7de !important;
#     margin: 16px 0 !important;
# }

# /* Info/success/error boxes */
# .stAlert {
#     border-radius: 6px !important;
#     border-left: 3px solid !important;
# }

# /* Download button */
# .stDownloadButton > button {
#     background: #0969da !important;
#     color: #ffffff !important;
#     border: 1px solid #0860ca !important;
#     border-radius: 6px !important;
#     font-family: 'Inter', sans-serif !important;
#     font-weight: 500 !important;
# }
# .stDownloadButton > button:hover {
#     background: #0860ca !important;
# }

# /* Expander */
# .streamlit-expanderHeader {
#     background: #ffffff !important;
#     border: 1px solid #d0d7de !important;
#     border-radius: 6px !important;
#     color: #1f2328 !important;
#     font-family: 'Inter', sans-serif !important;
# }

# /* Dataframe */
# [data-testid="stDataFrame"] {
#     border: 1px solid #d0d7de !important;
#     border-radius: 6px !important;
# }

# /* Spinner */
# .stSpinner > div {
#     border-top-color: #0969da !important;
# }

# /* Custom card styles */
# .doc-card {
#     background: #ffffff;
#     border: 1px solid #d0d7de;
#     border-radius: 8px;
#     padding: 20px;
#     margin-bottom: 12px;
#     transition: border-color 0.2s;
# }
# .doc-card:hover {
#     border-color: #636c76;
# }
# .section-title {
#     font-size: 13px;
#     font-weight: 600;
#     color: #636c76;
#     text-transform: uppercase;
#     letter-spacing: 0.08em;
#     margin-bottom: 12px;
#     padding-bottom: 8px;
#     border-bottom: 1px solid #d0d7de;
# }
# .repo-badge {
#     display: inline-flex;
#     align-items: center;
#     gap: 6px;
#     background: #ddf4ff;
#     border: 1px solid #54aeff;
#     border-radius: 20px;
#     padding: 4px 12px;
#     font-size: 13px;
#     font-weight: 500;
#     color: #0969da;
#     font-family: 'JetBrains Mono', monospace;
# }
# .status-badge {
#     display: inline-block;
#     padding: 2px 8px;
#     border-radius: 12px;
#     font-size: 11px;
#     font-weight: 600;
#     text-transform: uppercase;
#     letter-spacing: 0.05em;
# }
# .status-success { background: #dafbe1; color: #1a7f37; border: 1px solid #82cfff; }
# .status-running { background: #ddf4ff; color: #0969da; border: 1px solid #54aeff; }
# .status-failed  { background: #ffebe9; color: #cf222e; border: 1px solid #ff8182; }
# .status-queued  { background: #fff8c5; color: #9a6700; border: 1px solid #d4a72c; }

# .content-section {
#     background: #ffffff;
#     border: 1px solid #d0d7de;
#     border-radius: 8px;
#     padding: 24px;
#     margin-bottom: 16px;
# }
# .content-section h3 {
#     color: #1f2328;
#     font-size: 16px;
#     font-weight: 600;
#     margin-bottom: 16px;
#     padding-bottom: 10px;
#     border-bottom: 1px solid #d0d7de;
# }
# .commit-row {
#     display: flex;
#     align-items: center;
#     gap: 12px;
#     padding: 12px 0;
#     border-bottom: 1px solid #eaeef2;
# }
# .sha-badge {
#     background: #f6f8fa;
#     border: 1px solid #d0d7de;
#     border-radius: 4px;
#     padding: 2px 8px;
#     font-family: 'JetBrains Mono', monospace;
#     font-size: 12px;
#     color: #0969da;
# }
# .file-badge {
#     background: #ddf4ff;
#     border: 1px solid #54aeff;
#     border-radius: 4px;
#     padding: 1px 6px;
#     font-family: 'JetBrains Mono', monospace;
#     font-size: 11px;
#     color: #0969da;
# }
# .impact-high   { color: #cf222e; font-weight: 600; }
# .impact-medium { color: #9a6700; font-weight: 600; }
# .impact-low    { color: #1a7f37; font-weight: 600; }

# .hero-title {
#     font-size: 32px;
#     font-weight: 700;
#     color: #1f2328;
#     letter-spacing: -0.02em;
#     margin-bottom: 8px;
# }
# .hero-sub {
#     font-size: 16px;
#     color: #636c76;
#     margin-bottom: 32px;
# }
# .url-input-wrapper {
#     background: #ffffff;
#     border: 1px solid #d0d7de;
#     border-radius: 12px;
#     padding: 32px;
#     max-width: 640px;
#     margin: 0 auto;
# }
# .loading-step {
#     display: flex;
#     align-items: center;
#     gap: 10px;
#     padding: 8px 0;
#     font-size: 14px;
#     color: #636c76;
# }
# .loading-step.done { color: #1a7f37; }
# .loading-step.active { color: #0969da; }
# </style>
# """, unsafe_allow_html=True)

# # ── Config ────────────────────────────────────────────────────────────────────
# DOCS_FOLDER = os.getenv("DOCS_FOLDER", "./generated_docs")
# API_BASE    = "http://127.0.0.1:8001"
# os.makedirs(DOCS_FOLDER, exist_ok=True)

# if st_autorefresh:
#     st_autorefresh(interval=10_000, key="autorefresh")

# # ── Session state ─────────────────────────────────────────────────────────────
# if "view" not in st.session_state:
#     st.session_state.view = "home"          # home | dashboard | generating
# if "active_repo" not in st.session_state:
#     st.session_state.active_repo = None     # "owner/repo"
# if "active_branch" not in st.session_state:
#     st.session_state.active_branch = "main"
# if "active_job_id" not in st.session_state:
#     st.session_state.active_job_id = None
# if "doc_content" not in st.session_state:
#     st.session_state.doc_content = {}       # repo → parsed content
# if "left_tab" not in st.session_state:
#     st.session_state.left_tab = "design"    # design | commits

# # ── Helpers ───────────────────────────────────────────────────────────────────
# def parse_doc_filename(filename: str) -> dict:
#     init_pattern   = r"^design_([^_]+)_([^_]+)_([^_]+)_([0-9]{8})_([0-9]{6})\.docx$"
#     commit_pattern = r"^commit_([0-9a-fA-F]+)_([^_]+)_([0-9]{8})_([0-9]{6})\.docx$"
#     m = re.match(init_pattern, filename)
#     if m:
#         owner, repo, branch, dv, tv = m.groups()
#         dt = None
#         try: dt = datetime.strptime(f"{dv}{tv}", "%Y%m%d%H%M%S")
#         except: pass
#         return {"file": filename, "type": "initial", "owner": owner,
#                 "repo": repo, "branch": branch, "dt": dt}
#     m = re.match(commit_pattern, filename)
#     if m:
#         sha, branch, dv, tv = m.groups()
#         dt = None
#         try: dt = datetime.strptime(f"{dv}{tv}", "%Y%m%d%H%M%S")
#         except: pass
#         return {"file": filename, "type": "commit", "sha": sha,
#                 "branch": branch, "dt": dt}
#     return {"file": filename, "type": "unknown", "dt": None}

# def get_all_docs() -> list:
#     files = [f for f in os.listdir(DOCS_FOLDER) if f.endswith(".docx")]
#     parsed = [parse_doc_filename(f) for f in files]
#     return sorted(parsed, key=lambda x: x["dt"] or datetime.min, reverse=True)

# def get_repo_docs(owner: str, repo: str) -> dict:
#     all_docs = get_all_docs()
#     initial = [d for d in all_docs if d["type"] == "initial"
#                and d.get("owner") == owner and d.get("repo") == repo]
#     commits = [d for d in all_docs if d["type"] == "commit"]
#     return {"initial": initial, "commits": commits}

# def get_known_repos() -> list:
#     all_docs = get_all_docs()
#     seen = {}
#     for d in all_docs:
#         if d["type"] == "initial" and d.get("owner") and d.get("repo"):
#             key = f"{d['owner']}/{d['repo']}"
#             if key not in seen:
#                 seen[key] = {"owner": d["owner"], "repo": d["repo"],
#                              "branch": d.get("branch", "main"), "dt": d["dt"]}
#     return list(seen.values())

# def get_jobs(limit: int = 20) -> list:
#     try:
#         r = requests.get(f"{API_BASE}/jobs?limit={limit}", timeout=3)
#         if r.status_code == 200:
#             return r.json().get("jobs", [])
#     except: pass
#     return []

# def server_running() -> bool:
#     try:
#         r = requests.get(f"{API_BASE}/health", timeout=2)
#         return r.status_code == 200
#     except: return False

# def read_docx_sections(filepath: str) -> dict:
#     """
#     Parse a docx built by docx_builder.build_initial_docx() and return a dict:
#       { section_name: [content_lines], ... }

#     Strategy: walk body elements in order. Use the XML w:sz value on runs
#     to detect level-1 headings (28 half-points = 14pt) — this is the most
#     reliable signal because docx_builder._add_heading always sets it explicitly.
#     """
#     from docx import Document
#     from docx.oxml.ns import qn
#     from lxml import etree

#     SECTION_ORDER = [
#         "Project Overview",
#         "Architecture",
#         "Tools & Tech Stack",
#         "API Information",
#         "Code Quality Metrics",
#         "Data Flow",
#     ]

#     # Known level-1 heading texts → which section they open
#     HEADING_TO_SECTION = {
#         "document information":        "Project Overview",
#         "overview":                    "Project Overview",
#         "description":                 "Project Overview",
#         "system overview":             "Project Overview",
#         "purpose":                     "Project Overview",
#         "architecture":                "Architecture",
#         "architecture diagram":        "Architecture",
#         "tools & tech stack":          "Tools & Tech Stack",
#         "tools and tech stack":        "Tools & Tech Stack",
#         "configuration & tech stack":  "Tools & Tech Stack",
#         "code quality metrics":        "Code Quality Metrics",
#         "api endpoints":               "API Information",
#         "api information":             "API Information",
#         "data flow":                   "Data Flow",
#     }

#     sections = {s: [] for s in SECTION_ORDER}
#     current = "Project Overview"

#     doc = Document(filepath)
#     body = doc.element.body
#     para_list  = doc.paragraphs
#     table_list = doc.tables
#     para_idx  = 0
#     table_idx = 0

#     def get_run_sz(para) -> int:
#         """Return font size in half-points for the first visible-text run.
#         Tries XML w:sz first (most reliable), falls back to python-docx font.size API."""
#         # XML approach — checks runs that have actual text content
#         for r_el in para._p.findall(f".//{qn('w:r')}"):
#             t_el = r_el.find(qn("w:t"))
#             if t_el is None or not (t_el.text or "").strip():
#                 continue
#             rpr = r_el.find(qn("w:rPr"))
#             if rpr is not None:
#                 sz = rpr.find(qn("w:sz"))
#                 if sz is not None:
#                     val = sz.get(qn("w:val"))
#                     if val:
#                         try:
#                             return int(val)
#                         except Exception:
#                             pass
#         # python-docx API fallback
#         for run in para.runs:
#             if run.text.strip():
#                 try:
#                     if run.font.size:
#                         return int(run.font.size.pt * 2)
#                 except Exception:
#                     pass
#         return 0

#     def is_bold_para(para) -> bool:
#         text_runs = [r for r in para.runs if r.text.strip()]
#         return bool(text_runs) and all(r.bold for r in text_runs)

#     def is_heading_by_name(para) -> bool:
#         """Also catch Word built-in Heading styles as a safety net."""
#         style = para.style.name if para.style else ""
#         return "Heading 1" in style or "Heading 2" in style

#     def is_mono_para(para) -> bool:
#         return any(
#             r.font.name and r.font.name.lower() in
#             ("courier new", "courier", "consolas", "monospace")
#             for r in para.runs if r.text.strip()
#         )

#     def para_to_line(para) -> str:
#         style = para.style.name if para.style else ""
#         raw   = para.text.strip()
#         if not raw:
#             return ""
#         # Word built-in list styles
#         if "List Bullet" in style or style.startswith("List Bullet"):
#             return f"- {raw}"
#         if "List Number" in style or style.startswith("List Number"):
#             return f"1. {raw}"
#         # Monospace → code
#         if is_mono_para(para):
#             return f"`{raw}`"
#         # Bold sub-label
#         if is_bold_para(para):
#             return f"**{raw}**"
#         return raw

#     def table_to_lines(table) -> list:
#         rows = table.rows
#         if not rows:
#             return []
#         out = []
#         headers = [c.text.strip().replace("\n", " ") for c in rows[0].cells]
#         out.append("| " + " | ".join(headers) + " |")
#         out.append("| " + " | ".join(["---"] * len(headers)) + " |")
#         for row in rows[1:]:
#             cells = [c.text.strip().replace("\n", " ") for c in row.cells]
#             out.append("| " + " | ".join(cells) + " |")
#         return out

#     for child in body:
#         tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

#         if tag == "p":
#             if para_idx >= len(para_list):
#                 para_idx += 1
#                 continue
#             para = para_list[para_idx]
#             para_idx += 1
#             raw = para.text.strip()
#             if not raw:
#                 continue

#             # Detect level-1 heading: bold + sz >= 26 half-pts (13pt+)  OR Word Heading 1 style
#             sz = get_run_sz(para)
#             is_h1 = (is_bold_para(para) and sz >= 26) or ("Heading 1" in (para.style.name if para.style else ""))
#             is_h2 = (is_bold_para(para) and 20 <= sz < 26) or ("Heading 2" in (para.style.name if para.style else ""))

#             if is_h1:
#                 key = raw.lower().strip()
#                 # Try exact match first
#                 matched = HEADING_TO_SECTION.get(key)
#                 if not matched:
#                     # Try substring match
#                     for k, v in HEADING_TO_SECTION.items():
#                         if k in key or key in k:
#                             matched = v
#                             break
#                 if matched:
#                     current = matched
#                 else:
#                     # Unknown heading — treat as sub-label in current section
#                     sections[current].append(f"**{raw}**")
#                 continue

#             # Detect level-2 heading: bold + sz in [20,26) → sub-label
#             if is_h2:
#                 sections[current].append(f"**{raw}**")
#                 continue

#             line = para_to_line(para)
#             if line:
#                 sections[current].append(line)

#         elif tag == "tbl":
#             if table_idx >= len(table_list):
#                 table_idx += 1
#                 continue
#             tbl_lines = table_to_lines(table_list[table_idx])
#             table_idx += 1
#             if tbl_lines:
#                 sections[current].extend(tbl_lines)

#     return sections


# def read_docx_text(filepath: str) -> str:
#     """Legacy shim — returns joined text (used by commit doc parser)."""
#     try:
#         sections = read_docx_sections(filepath)
#         all_lines = []
#         for sec, lines in sections.items():
#             if lines:
#                 all_lines.append(f"## {sec}")
#                 all_lines.extend(lines)
#         return "\n".join(all_lines)
#     except Exception as e:
#         return f"Could not read document: {e}"

# # ── SIDEBAR ───────────────────────────────────────────────────────────────────
# with st.sidebar:
#     # Logo / Brand
#     st.markdown("""
#     <div style="padding: 8px 0 24px 0;">
#         <div style="font-size:22px; font-weight:700; color:#1f2328; letter-spacing:-0.02em;">
#             📄 DocSync
#         </div>
#         <div style="font-size:12px; color:#636c76; margin-top:2px;">
#             Auto-generated design docs
#         </div>
#     </div>
#     """, unsafe_allow_html=True)

#     # Server status dot
#     is_running = server_running()
#     status_color = "#1a7f37" if is_running else "#cf222e"
#     status_text  = "Server online" if is_running else "Server offline"
#     st.markdown(f"""
#     <div style="display:flex;align-items:center;gap:8px;padding:8px 12px;
#                 background:#f6f8fa;border:1px solid #d0d7de;border-radius:6px;
#                 margin-bottom:16px;">
#         <div style="width:8px;height:8px;border-radius:50%;background:{status_color};
#                     box-shadow:0 0 6px {status_color};"></div>
#         <span style="font-size:13px;color:#636c76;">{status_text}</span>
#     </div>
#     """, unsafe_allow_html=True)

#     # Nav: Home
#     if st.button("＋  New Repository", use_container_width=True):
#         st.session_state.view = "home"
#         st.rerun()

#     st.markdown('<div class="section-title" style="margin-top:20px;">YOUR REPOSITORIES</div>',
#                 unsafe_allow_html=True)

#     # List known repos
#     known = get_known_repos()
#     if not known:
#         st.markdown('<div style="font-size:13px;color:#636c76;padding:8px 0;">No repositories yet.<br>Add one above.</div>',
#                     unsafe_allow_html=True)
#     else:
#         for repo_info in known:
#             repo_key = f"{repo_info['owner']}/{repo_info['repo']}"
#             is_active = st.session_state.active_repo == repo_key
#             btn_style = "background:#f6f8fa;border:1px solid #0969da;" if is_active else ""
#             if st.button(f"📁  {repo_key}", key=f"repo_{repo_key}", use_container_width=True):
#                 st.session_state.active_repo  = repo_key
#                 st.session_state.active_branch = repo_info.get("branch", "main")
#                 st.session_state.view          = "dashboard"
#                 st.session_state.left_tab      = "design"
#                 st.rerun()

#     st.markdown("---")

#     # Recent jobs
#     st.markdown('<div class="section-title">RECENT JOBS</div>', unsafe_allow_html=True)
#     jobs = get_jobs(5)
#     if not jobs:
#         st.markdown('<div style="font-size:12px;color:#636c76;">No jobs yet.</div>',
#                     unsafe_allow_html=True)
#     for job in jobs[:5]:
#         status = job.get("status", "unknown")
#         icon = {"success": "✅", "running": "🔄", "failed": "❌", "queued": "⏳"}.get(status, "•")
#         sha  = job.get("commit_sha", "")[:10]
#         st.markdown(f"""
#         <div style="padding:6px 0;border-bottom:1px solid #eaeef2;">
#             <div style="font-size:12px;color:#636c76;">{icon} <span style="font-family:'JetBrains Mono',monospace;color:#0969da;">{sha}</span></div>
#             <div style="font-size:11px;color:#57606a;margin-top:2px;">{status}</div>
#         </div>
#         """, unsafe_allow_html=True)

#     if st.button("🔄 Refresh", use_container_width=True):
#         st.rerun()


# # ── MAIN CONTENT ──────────────────────────────────────────────────────────────

# # ════════════════════════════════════════════════════════════
# # VIEW: HOME — URL Input
# # ════════════════════════════════════════════════════════════
# if st.session_state.view == "home":
#     st.markdown("""
#     <div style="max-width:700px;margin:60px auto 0 auto;text-align:center;">
#         <div class="hero-title">Which repo would you like to document?</div>
#         <div class="hero-sub">Paste a GitHub URL to auto-generate a full system design document</div>
#     </div>
#     """, unsafe_allow_html=True)

#     col_l, col_c, col_r = st.columns([1, 2, 1])
#     with col_c:
#         github_url = st.text_input(
#             "",
#             placeholder="https://github.com/owner/repo",
#             label_visibility="collapsed",
#         )
#         branch = st.text_input("", value="main", placeholder="Branch (default: main)",
#                                label_visibility="collapsed")

#         generate_clicked = st.button("🚀  Generate Design Document", use_container_width=True, type="primary")

#         if generate_clicked:
#             if not github_url or "github.com" not in github_url:
#                 st.error("Please enter a valid GitHub URL")
#             elif not is_running:
#                 st.error("FastAPI server is not running. Start it with: `uvicorn main:app --port 8001`")
#             else:
#                 with st.spinner("Submitting..."):
#                     try:
#                         resp = requests.post(
#                             f"{API_BASE}/generate",
#                             json={"github_url": github_url, "branch": branch or "main"},
#                             timeout=10,
#                         )
#                         if resp.status_code == 200:
#                             data = resp.json()
#                             repo_key = f"{data['owner']}/{data['repo']}"
#                             st.session_state.active_repo   = repo_key
#                             st.session_state.active_branch = data.get("branch", "main")
#                             st.session_state.active_job_id = data.get("job_id")
#                             st.session_state.view          = "generating"
#                             st.rerun()
#                         else:
#                             st.error(f"Error: {resp.text}")
#                     except Exception as e:
#                         st.error(f"Could not reach server: {e}")

#     # Show existing repos as cards
#     known = get_known_repos()
#     if known:
#         st.markdown("---")
#         st.markdown('<div style="text-align:center;font-size:13px;color:#636c76;margin-bottom:16px;">OR OPEN AN EXISTING REPOSITORY</div>',
#                     unsafe_allow_html=True)
#         cols = st.columns(min(3, len(known)))
#         for i, repo_info in enumerate(known[:3]):
#             repo_key = f"{repo_info['owner']}/{repo_info['repo']}"
#             with cols[i % 3]:
#                 st.markdown(f"""
#                 <div class="doc-card" style="cursor:pointer;">
#                     <div style="font-size:14px;font-weight:600;color:#1f2328;margin-bottom:4px;">
#                         📁 {repo_key}
#                     </div>
#                     <div style="font-size:12px;color:#636c76;">Branch: {repo_info.get('branch','main')}</div>
#                     <div style="font-size:11px;color:#57606a;margin-top:4px;">
#                         {repo_info['dt'].strftime('%b %d, %Y') if repo_info.get('dt') else 'Unknown date'}
#                     </div>
#                 </div>
#                 """, unsafe_allow_html=True)
#                 if st.button(f"Open", key=f"open_{repo_key}", use_container_width=True):
#                     st.session_state.active_repo  = repo_key
#                     st.session_state.active_branch = repo_info.get("branch", "main")
#                     st.session_state.view          = "dashboard"
#                     st.session_state.left_tab      = "design"
#                     st.rerun()


# # ════════════════════════════════════════════════════════════
# # VIEW: GENERATING — Loading screen
# # ════════════════════════════════════════════════════════════
# elif st.session_state.view == "generating":
#     repo_key = st.session_state.active_repo or "repository"
#     job_id   = st.session_state.active_job_id

#     st.markdown(f"""
#     <div style="max-width:560px;margin:80px auto 0 auto;text-align:center;">
#         <div style="font-size:48px;margin-bottom:16px;">⚙️</div>
#         <div class="hero-title" style="font-size:24px;">Generating documentation</div>
#         <div class="hero-sub" style="font-size:14px;margin-bottom:32px;">
#             Scanning <span style="color:#0969da;font-family:'JetBrains Mono',monospace;">{repo_key}</span>
#             — this takes 2-5 minutes
#         </div>
#     </div>
#     """, unsafe_allow_html=True)

#     # Show live job status
#     if job_id:
#         try:
#             jr = requests.get(f"{API_BASE}/jobs/{job_id}", timeout=3)
#             if jr.status_code == 200:
#                 job = jr.json()
#                 status = job.get("status", "queued")
#                 steps_done = job.get("steps_completed", [])
#                 current = job.get("current_step", "")

#                 all_steps = [
#                     ("fetch_repo_context",        "Fetching repository files"),
#                     ("detect_repo_type",           "Detecting repo type"),
#                     ("analyze_tech_stack",         "Analyzing tech stack"),
#                     ("compute_metrics",            "Computing code metrics"),
#                     ("generate_repo_description",  "Generating system overview"),
#                     ("generate_architecture",      "Generating architecture section"),
#                     ("generate_api_section",       "Generating API documentation"),
#                     ("generate_data_flow",         "Generating data flow"),
#                     ("save_docx",                  "Saving design document (.docx)"),
#                     ("push_markdown",              "Pushing DOCUMENTATION.md to repo"),
#                     ("save_section_map",           "Saving section map"),
#                 ]

#                 col_l2, col_c2, col_r2 = st.columns([1, 2, 1])
#                 with col_c2:
#                     progress = len(steps_done) / max(len(all_steps), 1)
#                     st.progress(progress)
#                     for step_key, step_label in all_steps:
#                         if step_key in steps_done:
#                             st.markdown(f'<div class="loading-step done">✅ {step_label}</div>',
#                                         unsafe_allow_html=True)
#                         elif step_key == current:
#                             st.markdown(f'<div class="loading-step active">🔄 {step_label}...</div>',
#                                         unsafe_allow_html=True)
#                         else:
#                             st.markdown(f'<div class="loading-step">○ {step_label}</div>',
#                                         unsafe_allow_html=True)

#                     if status == "success":
#                         st.success("✅ Documentation generated successfully!")
#                         if st.button("📂 View Documentation", type="primary", use_container_width=True):
#                             st.session_state.view     = "dashboard"
#                             st.session_state.left_tab = "design"
#                             st.rerun()
#                     elif status == "failed":
#                         err = job.get("error", "Unknown error")
#                         st.error(f"❌ Generation failed: {err}")
#                         if st.button("← Try Again"):
#                             st.session_state.view = "home"
#                             st.rerun()
#         except Exception as e:
#             st.info(f"Checking job status... ({e})")
#     else:
#         col_l2, col_c2, col_r2 = st.columns([1, 2, 1])
#         with col_c2:
#             st.info("Job submitted. Waiting for updates...")


# # ════════════════════════════════════════════════════════════
# # VIEW: DASHBOARD — Two-panel layout
# # ════════════════════════════════════════════════════════════
# elif st.session_state.view == "dashboard":
#     repo_key = st.session_state.active_repo or ""
#     owner, repo_name = (repo_key.split("/", 1) + [""])[:2]
#     branch = st.session_state.active_branch

#     docs = get_repo_docs(owner, repo_name)
#     initial_docs = docs["initial"]
#     commit_docs  = docs["commits"]

#     # ── Top bar ───────────────────────────────────────────────────────────────
#     col_info, col_actions = st.columns([3, 1])
#     with col_info:
#         st.markdown(f"""
#         <div style="display:flex;align-items:center;gap:12px;padding:8px 0 16px 0;">
#             <div class="repo-badge">📁 {repo_key}</div>
#             <div style="font-size:12px;color:#636c76;padding:3px 8px;background:#ddf4ff;
#                         border:1px solid #d0d7de;border-radius:4px;">
#                 🌿 {branch}
#             </div>
#             <div style="font-size:12px;color:#636c76;">
#                 {len(initial_docs)} design doc{"s" if len(initial_docs)!=1 else ""} · {len(commit_docs)} commit doc{"s" if len(commit_docs)!=1 else ""}
#             </div>
#         </div>
#         """, unsafe_allow_html=True)
#     with col_actions:
#         github_url = f"https://github.com/{repo_key}/blob/{branch}/DOCUMENTATION.md"
#         st.markdown(f"""
#         <div style="text-align:right;padding-top:8px;">
#             <a href="{github_url}" target="_blank"
#                style="display:inline-flex;align-items:center;gap:6px;
#                       background:#ffffff;border:1px solid #d0d7de;border-radius:6px;
#                       padding:6px 14px;text-decoration:none;color:#1f2328;font-size:13px;
#                       font-weight:500;transition:border-color 0.2s;">
#                 🔗 View in GitHub
#             </a>
#         </div>
#         """, unsafe_allow_html=True)

#     st.markdown('<hr style="margin:0 0 16px 0;">', unsafe_allow_html=True)

#     # ── Two-column layout: LEFT nav | RIGHT content ───────────────────────────
#     left_col, right_col = st.columns([1, 3])

#     with left_col:
#         # Section nav
#         st.markdown('<div class="section-title">SECTIONS</div>', unsafe_allow_html=True)

#         design_btn = st.button("📋  Design Documentation",
#                                use_container_width=True,
#                                type="primary" if st.session_state.left_tab == "design" else "secondary")
#         commits_btn = st.button("🔀  Commit Changes & Impact",
#                                 use_container_width=True,
#                                 type="primary" if st.session_state.left_tab == "commits" else "secondary")

#         if design_btn:
#             st.session_state.left_tab = "design"
#             st.rerun()
#         if commits_btn:
#             st.session_state.left_tab = "commits"
#             st.rerun()

#         st.markdown("---")

#         # Design doc sub-sections
#         if st.session_state.left_tab == "design":
#             st.markdown('<div class="section-title">ON THIS PAGE</div>', unsafe_allow_html=True)
#             sections = [
#                 "Project Overview",
#                 "Architecture",
#                 "Tools & Tech Stack",
#                 "API Information",
#                 "Code Quality Metrics",
#                 "Data Flow",
#             ]
#             for s in sections:
#                 st.markdown(f'<div style="font-size:13px;color:#636c76;padding:4px 8px;'
#                             f'cursor:pointer;">• {s}</div>', unsafe_allow_html=True)

#         # Commit sub-sections
#         if st.session_state.left_tab == "commits":
#             st.markdown('<div class="section-title">ON THIS PAGE</div>', unsafe_allow_html=True)
#             for s in ["Latest Commits", "Impact Analysis", "Files Changed"]:
#                 st.markdown(f'<div style="font-size:13px;color:#636c76;padding:4px 8px;">• {s}</div>',
#                             unsafe_allow_html=True)

#     # ── RIGHT PANEL ───────────────────────────────────────────────────────────
#     with right_col:

#         # ── SECTION 1: DESIGN DOCUMENTATION ──────────────────────────────────
#         if st.session_state.left_tab == "design":
#             st.markdown("""
#             <div style="font-size:24px;font-weight:700;color:#1f2328;
#                         letter-spacing:-0.02em;margin-bottom:4px;">
#                 📋 Design Documentation
#             </div>
#             <div style="font-size:14px;color:#636c76;margin-bottom:24px;">
#                 Auto-generated system design document for this repository
#             </div>
#             """, unsafe_allow_html=True)

#             if not initial_docs:
#                 st.info("No design document generated yet. The document will appear here after generation completes.")
#             else:
#                 latest = initial_docs[0]
#                 filepath = os.path.join(DOCS_FOLDER, latest["file"])

#                 # Doc metadata banner
#                 st.markdown(f"""
#                 <div style="background:#ddf4ff;border:1px solid #0969da;border-radius:8px;
#                             padding:16px 20px;margin-bottom:20px;display:flex;
#                             align-items:center;justify-content:space-between;">
#                     <div>
#                         <div style="font-size:13px;color:#0969da;font-weight:500;margin-bottom:2px;">
#                             Latest Design Document
#                         </div>
#                         <div style="font-size:12px;color:#636c76;">
#                             Generated: {latest['dt'].strftime('%B %d, %Y at %H:%M') if latest.get('dt') else 'Unknown'}
#                             &nbsp;·&nbsp; Branch: {latest.get('branch','main')}
#                         </div>
#                     </div>
#                     <div style="font-size:11px;color:#1a7f37;background:#dafbe1;
#                                 border:1px solid #1a7f37;border-radius:12px;padding:3px 10px;">
#                         ✓ Up to date
#                     </div>
#                 </div>
#                 """, unsafe_allow_html=True)

#                 # ── Parse docx directly into sections ────────────────────────
#                 SECTION_ORDER = [
#                     "Project Overview",
#                     "Architecture",
#                     "Tools & Tech Stack",
#                     "API Information",
#                     "Code Quality Metrics",
#                     "Data Flow",
#                 ]
#                 try:
#                     sections_map = read_docx_sections(filepath)
#                 except Exception as _e:
#                     st.error(f"Could not parse document: {_e}")
#                     sections_map = {s: [] for s in SECTION_ORDER}


#                 # ── Display each section ──────────────────────────────────────
#                 icons = {
#                     "Project Overview": "🏠",
#                     "Architecture":     "🏗️",
#                     "Tools & Tech Stack": "🛠️",
#                     "API Information":  "🔌",
#                     "Code Quality Metrics": "📊",
#                     "Data Flow":        "🔄",
#                 }

#                 def render_section_content(sec_lines: list):
#                     """Render section lines: sub-headings on own line, bullets as lists, tables as dataframes."""
#                     pending_text = []
#                     pending_table = []
#                     in_table = False

#                     def flush_text():
#                         if pending_text:
#                             block = "\n".join(pending_text).strip()
#                             if block:
#                                 st.markdown(block)
#                             pending_text.clear()

#                     def flush_table():
#                         if not pending_table:
#                             return
#                         rows = []
#                         for tl in pending_table:
#                             tl_s = tl.strip()
#                             if re.match(r"^\|[-| :]+\|$", tl_s):
#                                 continue
#                             if tl_s.startswith("|") and tl_s.endswith("|"):
#                                 parts = [p.strip() for p in tl_s.strip("|").split("|")]
#                                 if parts and any(p for p in parts):
#                                     rows.append(parts)
#                         if len(rows) > 1:
#                             try:
#                                 import pandas as pd
#                                 df = pd.DataFrame(rows[1:], columns=rows[0])
#                                 st.dataframe(df, use_container_width=True, hide_index=True)
#                             except Exception:
#                                 st.markdown("\n".join(pending_table))
#                         elif rows:
#                             st.markdown("\n".join(pending_table))
#                         pending_table.clear()

#                     for line in sec_lines:
#                         stripped = line.strip()

#                         # Sub-heading line — bold label on its own (e.g. **Overview**, **Description**)
#                         is_subheading = (
#                             stripped.startswith("**") and stripped.endswith("**")
#                             and "\n" not in stripped
#                             and not stripped.startswith("**`")  # not a bullet bold label
#                             and len(stripped) < 80             # headings are short
#                             and not stripped[2:-2].strip().startswith("-")
#                         )

#                         # Real markdown table line
#                         is_table_line = (
#                             stripped.startswith("|") and stripped.endswith("|")
#                             and stripped.count("|") >= 3
#                         )

#                         if is_subheading:
#                             # Flush whatever came before
#                             flush_text()
#                             flush_table()
#                             in_table = False
#                             # Render sub-heading as its own styled element
#                             label = stripped[2:-2].strip()
#                             st.markdown(
#                                 f'<div style="font-size:13px;font-weight:600;'
#                                 f'color:var(--color-text-primary,#1f2328);'
#                                 f'margin:14px 0 6px 0;padding-bottom:4px;'
#                                 f'border-bottom:1px solid #d0d7de;">{label}</div>',
#                                 unsafe_allow_html=True
#                             )
#                         elif is_table_line:
#                             if not in_table:
#                                 flush_text()
#                                 in_table = True
#                             pending_table.append(line)
#                         else:
#                             if in_table:
#                                 flush_table()
#                                 in_table = False
#                             # Skip blank lines between sub-heading and first bullet
#                             if stripped or pending_text:
#                                 pending_text.append(line)

#                     flush_text()
#                     flush_table()

#                 for sec_name in SECTION_ORDER:
#                     sec_lines = sections_map[sec_name]
#                     if not any(l.strip() for l in sec_lines):
#                         continue
#                     with st.expander(
#                         f"{icons.get(sec_name, '•')} {sec_name}",
#                         expanded=(sec_name == "Project Overview")
#                     ):
#                         render_section_content(sec_lines)

#                 # Download button
#                 st.markdown("---")
#                 col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])
#                 with col_dl1:
#                     if os.path.exists(filepath):
#                         with open(filepath, "rb") as fh:
#                             st.download_button(
#                                 label="⬇️  Download Design Doc (.docx)",
#                                 data=fh.read(),
#                                 file_name=latest["file"],
#                                 key=f"dl_design_{latest['file']}",
#                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                                 use_container_width=True,
#                             )
#                     else:
#                         st.warning("Document file not found.")
#                 with col_dl2:
#                     st.markdown(f"""
#                     <a href="https://github.com/{repo_key}/blob/{branch}/DOCUMENTATION.md"
#                        target="_blank"
#                        style="display:block;text-align:center;background:#ffffff;
#                               border:1px solid #d0d7de;border-radius:6px;padding:8px;
#                               text-decoration:none;color:#1f2328;font-size:14px;">
#                         🔗 View in GitHub Repo
#                     </a>
#                     """, unsafe_allow_html=True)

#                 # Older versions
#                 if len(initial_docs) > 1:
#                     with st.expander(f"📁 Older versions ({len(initial_docs)-1})"):
#                         for old_doc in initial_docs[1:]:
#                             old_path = os.path.join(DOCS_FOLDER, old_doc["file"])
#                             c1, c2 = st.columns([3, 1])
#                             with c1:
#                                 dt_str = old_doc['dt'].strftime('%b %d, %Y %H:%M') if old_doc.get('dt') else 'Unknown'
#                                 st.markdown(f'<div style="font-size:13px;color:#636c76;">{dt_str}</div>',
#                                             unsafe_allow_html=True)
#                             with c2:
#                                 if os.path.exists(old_path):
#                                     with open(old_path, "rb") as fh:
#                                         st.download_button("⬇", data=fh.read(),
#                                                            file_name=old_doc["file"],
#                                                            key=f"old_{old_doc['file']}",
#                                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

#         # ── SECTION 2: COMMIT CHANGES & IMPACT ANALYSIS ───────────────────────
#         elif st.session_state.left_tab == "commits":
#             st.markdown("""
#             <div style="font-size:24px;font-weight:700;color:#1f2328;
#                         letter-spacing:-0.02em;margin-bottom:4px;">
#                 🔀 Commit Changes & Impact Analysis
#             </div>
#             <div style="font-size:14px;color:#636c76;margin-bottom:24px;">
#                 Auto-triggered on every GitHub push — shows what changed and its impact
#             </div>
#             """, unsafe_allow_html=True)

#             if not commit_docs:
#                 st.markdown("""
#                 <div class="content-section" style="text-align:center;padding:48px;">
#                     <div style="font-size:32px;margin-bottom:12px;">🔗</div>
#                     <div style="font-size:16px;color:#1f2328;font-weight:500;margin-bottom:8px;">
#                         No commit updates yet
#                     </div>
#                     <div style="font-size:14px;color:#636c76;">
#                         Set up a GitHub webhook pointing to<br>
#                         <span style="font-family:'JetBrains Mono',monospace;color:#0969da;">
#                             POST http://your-server:8001/webhook
#                         </span><br><br>
#                         Every push will automatically generate a commit analysis report here.
#                     </div>
#                 </div>
#                 """, unsafe_allow_html=True)
#             else:
#                 # Show each commit doc
#                 for i, commit_doc in enumerate(commit_docs[:10]):
#                     sha = commit_doc.get("sha", "unknown")
#                     branch_c = commit_doc.get("branch", "main")
#                     dt_str = commit_doc["dt"].strftime("%b %d, %Y at %H:%M") if commit_doc.get("dt") else "Unknown"
#                     filepath = os.path.join(DOCS_FOLDER, commit_doc["file"])

#                     with st.expander(
#                         f"🔀 Commit `{sha[:7]}` — {dt_str}",
#                         expanded=(i == 0)
#                     ):
#                         # Commit header
#                         st.markdown(f"""
#                         <div style="display:flex;align-items:center;gap:12px;
#                                     padding:12px;background:#f6f8fa;border-radius:6px;
#                                     margin-bottom:16px;">
#                             <div class="sha-badge">{sha[:7]}</div>
#                             <div style="font-size:12px;color:#636c76;">
#                                 🌿 {branch_c} &nbsp;·&nbsp; 📅 {dt_str}
#                             </div>
#                         </div>
#                         """, unsafe_allow_html=True)

#                         if os.path.exists(filepath):
#                             doc_text = read_docx_text(filepath)

#                             # Parse tables from doc
#                             lines = doc_text.split("\n")

#                             # Extract sections
#                             sections_updated = []
#                             commit_table_lines = []
#                             impact_table_lines = []
#                             in_commit = False
#                             in_impact = False

#                             for line in lines:
#                                 ll = line.lower()
#                                 if "sections updated" in ll:
#                                     in_commit = False; in_impact = False
#                                 elif "commit change" in ll:
#                                     in_commit = True; in_impact = False
#                                 elif "impact analysis" in ll:
#                                     in_commit = False; in_impact = True
#                                 elif in_commit and "|" in line:
#                                     commit_table_lines.append(line)
#                                 elif in_impact and "|" in line:
#                                     impact_table_lines.append(line)
#                                 elif "updated:" in ll or "unchanged:" in ll:
#                                     sections_updated.append(line.strip())

#                             # Sections updated badges
#                             if sections_updated:
#                                 st.markdown("**Sections Updated**")
#                                 badge_html = ""
#                                 for s in sections_updated:
#                                     color = "#1a4731" if "updated" in s.lower() else "#1c1c1c"
#                                     text_color = "#1a7f37" if "updated" in s.lower() else "#57606a"
#                                     badge_html += f'<span style="background:{color};color:{text_color};border-radius:4px;padding:2px 8px;font-size:12px;margin-right:6px;">{s}</span>'
#                                 st.markdown(badge_html, unsafe_allow_html=True)
#                                 st.markdown("")

#                             # Commit change details table
#                             if commit_table_lines:
#                                 st.markdown("**📝 Files Changed**")
#                                 import pandas as pd
#                                 rows = []
#                                 for row_line in commit_table_lines:
#                                     if "|" in row_line and not re.match(r"^[-| :]+$", row_line):
#                                         parts = [p.strip() for p in row_line.strip("|").split("|")]
#                                         if parts and any(p for p in parts):
#                                             rows.append(parts)
#                                 if len(rows) > 1:
#                                     try:
#                                         df = pd.DataFrame(rows[1:], columns=rows[0])
#                                         st.dataframe(df, use_container_width=True, hide_index=True)
#                                     except:
#                                         st.code("\n".join(commit_table_lines))
#                                 st.markdown("")

#                             # Impact analysis table
#                             if impact_table_lines:
#                                 st.markdown("**⚡ Impact Analysis**")
#                                 rows = []
#                                 for row_line in impact_table_lines:
#                                     if "|" in row_line and not re.match(r"^[-| :]+$", row_line):
#                                         parts = [p.strip() for p in row_line.strip("|").split("|")]
#                                         if parts and any(p for p in parts):
#                                             rows.append(parts)
#                                 if len(rows) > 1:
#                                     try:
#                                         import pandas as pd
#                                         df = pd.DataFrame(rows[1:], columns=rows[0])
#                                         st.dataframe(df, use_container_width=True, hide_index=True)
#                                     except:
#                                         st.code("\n".join(impact_table_lines))
#                             else:
#                                 # Show raw text if no tables found
#                                 preview = "\n".join(lines[:30])
#                                 if preview.strip():
#                                     st.text(preview)

#                             # Download this commit doc
#                             st.markdown("---")
#                             with open(filepath, "rb") as fh:
#                                 st.download_button(
#                                     label=f"⬇️ Download Commit Report — {sha[:7]}",
#                                     data=fh.read(),
#                                     file_name=commit_doc["file"],
#                                     key=f"dl_commit_{sha}_{i}",
#                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                                 )
#                         else:
#                             st.warning("Document file not found on disk.")

#                 # Webhook setup reminder
#                 st.markdown("""
#                 <div style="background:#ffffff;border:1px solid #d0d7de;border-radius:8px;
#                             padding:16px;margin-top:16px;">
#                     <div style="font-size:13px;font-weight:600;color:#1f2328;margin-bottom:8px;">
#                         🔗 Webhook Setup
#                     </div>
#                     <div style="font-size:12px;color:#636c76;">
#                         New commits are detected automatically via GitHub webhook.<br>
#                         Payload URL: <span style="font-family:'JetBrains Mono',monospace;color:#0969da;">
#                         http://your-server:8001/webhook</span>
#                     </div>
#                 </div>
#                 """, unsafe_allow_html=True)
import os
import sys
 
# Add project root directory to Python path FIRST
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
import env_loader  # noqa: F401
import re
import sys
import threading
import time
import json
from datetime import datetime
from pathlib import Path

import requests
import streamlit as st

try:
    from streamlit_autorefresh import st_autorefresh
except (ImportError, ModuleNotFoundError):
    st_autorefresh = None

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DocSync",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

/* Global */
html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* Hide default streamlit elements */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
.stDeployButton {display: none;}

/* Main background */
.stApp {
    background: #f6f8fa;
    color: #1f2328;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid #d0d7de !important;
}
[data-testid="stSidebar"] * {
    color: #1f2328 !important;
}

/* Buttons */
.stButton > button {
    background: #1f883d !important;
    color: #ffffff !important;
    border: 1px solid #1a7f37 !important;
    border-radius: 6px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 500 !important;
    padding: 8px 16px !important;
    transition: all 0.2s !important;
}
.stButton > button:hover {
    background: #1a7f37 !important;
    border-color: #157832 !important;
}

/* Input fields */
.stTextInput > div > div > input {
    background: #ffffff !important;
    border: 1px solid #d0d7de !important;
    color: #1f2328 !important;
    border-radius: 6px !important;
    font-family: 'Inter', sans-serif !important;
}
.stTextInput > div > div > input:focus {
    border-color: #0969da !important;
    box-shadow: 0 0 0 3px rgba(9,105,218,0.12) !important;
}

/* Text area */
.stTextArea textarea {
    background: #ffffff !important;
    border: 1px solid #d0d7de !important;
    color: #1f2328 !important;
    border-radius: 6px !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 13px !important;
}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    background: transparent !important;
    border-bottom: 1px solid #d0d7de !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    color: #636c76 !important;
    border: none !important;
    padding: 8px 16px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    font-weight: 500 !important;
}
.stTabs [aria-selected="true"] {
    color: #1f2328 !important;
    border-bottom: 2px solid #fd8c73 !important;
}

/* Metric cards */
[data-testid="stMetric"] {
    background: #ffffff !important;
    border: 1px solid #d0d7de !important;
    border-radius: 8px !important;
    padding: 16px !important;
}
[data-testid="stMetricLabel"] {
    color: #636c76 !important;
    font-size: 12px !important;
    font-weight: 500 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
}
[data-testid="stMetricValue"] {
    color: #1f2328 !important;
    font-size: 24px !important;
    font-weight: 700 !important;
}

/* Divider */
hr {
    border-color: #d0d7de !important;
    margin: 16px 0 !important;
}

/* Info/success/error boxes */
.stAlert {
    border-radius: 6px !important;
    border-left: 3px solid !important;
}

/* Download button */
.stDownloadButton > button {
    background: #0969da !important;
    color: #ffffff !important;
    border: 1px solid #0860ca !important;
    border-radius: 6px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 500 !important;
}
.stDownloadButton > button:hover {
    background: #0860ca !important;
}

/* Expander */
.streamlit-expanderHeader {
    background: #ffffff !important;
    border: 1px solid #d0d7de !important;
    border-radius: 6px !important;
    color: #1f2328 !important;
    font-family: 'Inter', sans-serif !important;
}

/* Dataframe */
[data-testid="stDataFrame"] {
    border: 1px solid #d0d7de !important;
    border-radius: 6px !important;
}

/* Spinner */
.stSpinner > div {
    border-top-color: #0969da !important;
}

/* Custom card styles */
.doc-card {
    background: #ffffff;
    border: 1px solid #d0d7de;
    border-radius: 8px;
    padding: 20px;
    margin-bottom: 12px;
    transition: border-color 0.2s;
}
.doc-card:hover {
    border-color: #636c76;
}
.section-title {
    font-size: 13px;
    font-weight: 600;
    color: #636c76;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    margin-bottom: 12px;
    padding-bottom: 8px;
    border-bottom: 1px solid #d0d7de;
}
.repo-badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: #ddf4ff;
    border: 1px solid #54aeff;
    border-radius: 20px;
    padding: 4px 12px;
    font-size: 13px;
    font-weight: 500;
    color: #0969da;
    font-family: 'JetBrains Mono', monospace;
}
.status-badge {
    display: inline-block;
    padding: 2px 8px;
    border-radius: 12px;
    font-size: 11px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.05em;
}
.status-success { background: #dafbe1; color: #1a7f37; border: 1px solid #82cfff; }
.status-running { background: #ddf4ff; color: #0969da; border: 1px solid #54aeff; }
.status-failed  { background: #ffebe9; color: #cf222e; border: 1px solid #ff8182; }
.status-queued  { background: #fff8c5; color: #9a6700; border: 1px solid #d4a72c; }

.content-section {
    background: #ffffff;
    border: 1px solid #d0d7de;
    border-radius: 8px;
    padding: 24px;
    margin-bottom: 16px;
}
.content-section h3 {
    color: #1f2328;
    font-size: 16px;
    font-weight: 600;
    margin-bottom: 16px;
    padding-bottom: 10px;
    border-bottom: 1px solid #d0d7de;
}
.commit-row {
    display: flex;
    align-items: center;
    gap: 12px;
    padding: 12px 0;
    border-bottom: 1px solid #eaeef2;
}
.sha-badge {
    background: #f6f8fa;
    border: 1px solid #d0d7de;
    border-radius: 4px;
    padding: 2px 8px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    color: #0969da;
}
.file-badge {
    background: #ddf4ff;
    border: 1px solid #54aeff;
    border-radius: 4px;
    padding: 1px 6px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    color: #0969da;
}
.impact-high   { color: #cf222e; font-weight: 600; }
.impact-medium { color: #9a6700; font-weight: 600; }
.impact-low    { color: #1a7f37; font-weight: 600; }

.hero-title {
    font-size: 32px;
    font-weight: 700;
    color: #1f2328;
    letter-spacing: -0.02em;
    margin-bottom: 8px;
}
.hero-sub {
    font-size: 16px;
    color: #636c76;
    margin-bottom: 32px;
}
.url-input-wrapper {
    background: #ffffff;
    border: 1px solid #d0d7de;
    border-radius: 12px;
    padding: 32px;
    max-width: 640px;
    margin: 0 auto;
}
.loading-step {
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 8px 0;
    font-size: 14px;
    color: #636c76;
}
.loading-step.done { color: #1a7f37; }
.loading-step.active { color: #0969da; }
</style>
""", unsafe_allow_html=True)

# ── Config ────────────────────────────────────────────────────────────────────
DOCS_FOLDER = os.getenv("DOCS_FOLDER", "./generated_docs")
API_BASE    = "http://127.0.0.1:8001"
os.makedirs(DOCS_FOLDER, exist_ok=True)

if st_autorefresh:
    st_autorefresh(interval=10_000, key="autorefresh")

# ── Session state ─────────────────────────────────────────────────────────────
if "view" not in st.session_state:
    st.session_state.view = "home"          # home | dashboard | generating
if "active_repo" not in st.session_state:
    st.session_state.active_repo = None     # "owner/repo"
if "active_branch" not in st.session_state:
    st.session_state.active_branch = "main"
if "active_job_id" not in st.session_state:
    st.session_state.active_job_id = None
if "doc_content" not in st.session_state:
    st.session_state.doc_content = {}       # repo → parsed content
if "left_tab" not in st.session_state:
    st.session_state.left_tab = "design"    # design | commits

# ── Helpers ───────────────────────────────────────────────────────────────────
def parse_doc_filename(filename: str) -> dict:
    init_pattern   = r"^design_([^_]+)_([^_]+)_([^_]+)_([0-9]{8})_([0-9]{6})\.docx$"
    commit_pattern = r"^commit_([0-9a-fA-F]+)_([^_]+)_([0-9]{8})_([0-9]{6})\.docx$"
    m = re.match(init_pattern, filename)
    if m:
        owner, repo, branch, dv, tv = m.groups()
        dt = None
        try: dt = datetime.strptime(f"{dv}{tv}", "%Y%m%d%H%M%S")
        except: pass
        return {"file": filename, "type": "initial", "owner": owner,
                "repo": repo, "branch": branch, "dt": dt}
    m = re.match(commit_pattern, filename)
    if m:
        sha, branch, dv, tv = m.groups()
        dt = None
        try: dt = datetime.strptime(f"{dv}{tv}", "%Y%m%d%H%M%S")
        except: pass
        return {"file": filename, "type": "commit", "sha": sha,
                "branch": branch, "dt": dt}
    return {"file": filename, "type": "unknown", "dt": None}

def get_all_docs() -> list:
    files = [f for f in os.listdir(DOCS_FOLDER) if f.endswith(".docx")]
    parsed = [parse_doc_filename(f) for f in files]
    return sorted(parsed, key=lambda x: x["dt"] or datetime.min, reverse=True)

def get_repo_docs(owner: str, repo: str) -> dict:
    all_docs = get_all_docs()
    initial = [d for d in all_docs if d["type"] == "initial"
               and d.get("owner") == owner and d.get("repo") == repo]
    commits = [d for d in all_docs if d["type"] == "commit"]
    return {"initial": initial, "commits": commits}

def get_known_repos() -> list:
    all_docs = get_all_docs()
    seen = {}
    for d in all_docs:
        if d["type"] == "initial" and d.get("owner") and d.get("repo"):
            key = f"{d['owner']}/{d['repo']}"
            if key not in seen:
                seen[key] = {"owner": d["owner"], "repo": d["repo"],
                             "branch": d.get("branch", "main"), "dt": d["dt"]}
    return list(seen.values())

def get_jobs(limit: int = 20) -> list:
    try:
        r = requests.get(f"{API_BASE}/jobs?limit={limit}", timeout=3)
        if r.status_code == 200:
            return r.json().get("jobs", [])
    except: pass
    return []

def server_running() -> bool:
    try:
        r = requests.get(f"{API_BASE}/health", timeout=2)
        return r.status_code == 200
    except: return False

def read_docx_sections(filepath: str) -> dict:
    """
    Parse a docx built by docx_builder.build_initial_docx() and return a dict:
      { section_name: [content_lines], ... }

    Strategy: walk body elements in order. Use the XML w:sz value on runs
    to detect level-1 headings (28 half-points = 14pt) — this is the most
    reliable signal because docx_builder._add_heading always sets it explicitly.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    SECTION_ORDER = [
        "Project Overview",
        "Architecture",
        "Tools & Tech Stack",
        "API Information",
        "Code Quality Metrics",
        "Data Flow",
    ]

    # Exact heading texts → target section.
    # ORDER MATTERS: more specific keys must come before shorter keys that are substrings of them.
    # e.g. "architecture diagram" before "architecture", "data flow description" before "description"
    HEADING_TO_SECTION = {
        # Project Overview (must come before short keys like "description")
        "document information":           "Project Overview",
        "system overview":                "Project Overview",
        "codebase structure":             "Project Overview",
        "what the codebase does":         "Project Overview",
        "overview":                       "Project Overview",
        "purpose":                        "Project Overview",
        # Architecture (must come before "architecture" to catch longer variants)
        # NOTE: "architecture diagram" is intentionally NOT here — it's in SUBLABEL_EXACT
        # because it's embedded inside repo_description, not a real section boundary.
        "architecture (updated)":         "Architecture",
        "architecture":                   "Architecture",
        # Tools & Tech Stack
        "tools & tech stack":             "Tools & Tech Stack",
        "tools and tech stack":           "Tools & Tech Stack",
        "configuration & tech stack":     "Tools & Tech Stack",
        # Code Quality
        "code quality metrics":           "Code Quality Metrics",
        "code quality metrics (updated)": "Code Quality Metrics",
        # API
        "api endpoints (updated)":        "API Information",
        "api endpoints":                  "API Information",
        "api information":                "API Information",
        # Data Flow (must come before "description" to catch "data flow description")
        "data flow (updated)":            "Data Flow",
        "data flow description":          "Data Flow",
        "data flow":                      "Data Flow",
        # Catch-all for "description" — only matches if nothing above matched
        "description":                    "Project Overview",
    }

    # Sub-headings that should stay as labels inside their current section (not switch)
    SUBLABEL_EXACT = {
        "high-level design", "key components", "component interactions", "entry points",
        "languages", "frameworks & libraries", "infrastructure",
        "data models", "storage",
        "work orders", "engineers", "parts", "documents", "authentication", "error handling",
        # "Architecture Diagram" is embedded inside repo_description LLM output —
        # it's a diagram label within Project Overview, not the Architecture section heading.
        "architecture diagram",
    }

    def match_heading(raw_lower: str):
        """Return target section name for a level-1 heading, or None if it's a sub-label."""
        # First check if it's a known sub-label — don't switch section for these
        if raw_lower in SUBLABEL_EXACT:
            return None
        # Exact match
        if raw_lower in HEADING_TO_SECTION:
            return HEADING_TO_SECTION[raw_lower]
        # Substring match — iterate in insertion order (specific → general)
        for k, v in HEADING_TO_SECTION.items():
            if k in raw_lower:   # k is a substring of the heading
                return v
        return None  # Unknown heading → sub-label

    sections = {s: [] for s in SECTION_ORDER}
    current = "Project Overview"

    doc = Document(filepath)
    body = doc.element.body
    para_list  = doc.paragraphs
    table_list = doc.tables
    para_idx  = 0
    table_idx = 0

    def get_run_sz(para) -> int:
        """Return font size in half-points for the first visible-text run.
        Tries XML w:sz first (most reliable), falls back to python-docx font.size API."""
        # XML approach — checks runs that have actual text content
        for r_el in para._p.findall(f".//{qn('w:r')}"):
            t_el = r_el.find(qn("w:t"))
            if t_el is None or not (t_el.text or "").strip():
                continue
            rpr = r_el.find(qn("w:rPr"))
            if rpr is not None:
                sz = rpr.find(qn("w:sz"))
                if sz is not None:
                    val = sz.get(qn("w:val"))
                    if val:
                        try:
                            return int(val)
                        except Exception:
                            pass
        # python-docx API fallback
        for run in para.runs:
            if run.text.strip():
                try:
                    if run.font.size:
                        return int(run.font.size.pt * 2)
                except Exception:
                    pass
        return 0

    def is_bold_para(para) -> bool:
        text_runs = [r for r in para.runs if r.text.strip()]
        return bool(text_runs) and all(r.bold for r in text_runs)

    def is_heading_by_name(para) -> bool:
        """Also catch Word built-in Heading styles as a safety net."""
        style = para.style.name if para.style else ""
        return "Heading 1" in style or "Heading 2" in style

    def is_mono_para(para) -> bool:
        return any(
            r.font.name and r.font.name.lower() in
            ("courier new", "courier", "consolas", "monospace")
            for r in para.runs if r.text.strip()
        )

    def para_to_line(para) -> str:
        style = para.style.name if para.style else ""
        raw   = para.text.strip()
        if not raw:
            return ""
        # Word built-in list styles
        if "List Bullet" in style or style.startswith("List Bullet"):
            return f"- {raw}"
        if "List Number" in style or style.startswith("List Number"):
            return f"1. {raw}"
        # Monospace → render as code (catches mermaid lines from _render_mermaid_block)
        if is_mono_para(para):
            return f"`{raw}`"
        # Bold sub-label (catches "Architecture Diagram" label at sz=18 from _render_mermaid_block)
        if is_bold_para(para):
            return f"**{raw}**"
        return raw

    def is_mermaid_line(para) -> bool:
        """Detect mermaid code lines: Courier New at ~8pt (sz=16) with grey background shading."""
        if not is_mono_para(para):
            return False
        # Check for grey background shading (added by _render_mermaid_block)
        tc_pr = para._p.find(f".//{qn('w:pPr')}")
        if tc_pr is not None:
            shd = tc_pr.find(qn("w:shd"))
            if shd is not None:
                fill = shd.get(qn("w:fill"), "").upper()
                if fill in ("F2F2F2", "EEEEEE", "E8E8E8"):
                    return True
        return is_mono_para(para)

    def table_to_lines(table) -> list:
        rows = table.rows
        if not rows:
            return []
        out = []
        headers = [c.text.strip().replace("\n", " ") for c in rows[0].cells]
        out.append("| " + " | ".join(headers) + " |")
        out.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows[1:]:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            out.append("| " + " | ".join(cells) + " |")
        return out

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_idx >= len(para_list):
                para_idx += 1
                continue
            para = para_list[para_idx]
            para_idx += 1
            raw = para.text.strip()
            if not raw:
                continue

            # Detect level-1 heading: bold + sz >= 26 half-pts (13pt+)  OR Word Heading 1 style
            sz = get_run_sz(para)
            is_h1 = (is_bold_para(para) and sz >= 26) or ("Heading 1" in (para.style.name if para.style else ""))
            is_h2 = (is_bold_para(para) and 20 <= sz < 26) or ("Heading 2" in (para.style.name if para.style else ""))

            if is_h1:
                key = raw.lower().strip()
                matched = match_heading(key)
                if matched:
                    current = matched
                else:
                    # Unknown heading or known sub-label → render as bold label in current section
                    sections[current].append(f"**{raw}**")
                continue

            # Detect level-2 heading: bold + sz in [20,26) → always a sub-label
            if is_h2:
                sections[current].append(f"**{raw}**")
                continue

            line = para_to_line(para)
            if line:
                sections[current].append(line)

        elif tag == "tbl":
            if table_idx >= len(table_list):
                table_idx += 1
                continue
            tbl_lines = table_to_lines(table_list[table_idx])
            table_idx += 1
            if tbl_lines:
                sections[current].extend(tbl_lines)

    return sections


def read_docx_text(filepath: str) -> str:
    """Legacy shim — returns joined text (used by commit doc parser)."""
    try:
        sections = read_docx_sections(filepath)
        all_lines = []
        for sec, lines in sections.items():
            if lines:
                all_lines.append(f"## {sec}")
                all_lines.extend(lines)
        return "\n".join(all_lines)
    except Exception as e:
        return f"Could not read document: {e}"

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    # Logo / Brand
    st.markdown("""
    <div style="padding: 8px 0 24px 0;">
        <div style="font-size:22px; font-weight:700; color:#1f2328; letter-spacing:-0.02em;">
            📄 DocSync
        </div>
        <div style="font-size:12px; color:#636c76; margin-top:2px;">
            Auto-generated design docs
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Server status dot
    is_running = server_running()
    status_color = "#1a7f37" if is_running else "#cf222e"
    status_text  = "Server online" if is_running else "Server offline"
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:8px;padding:8px 12px;
                background:#f6f8fa;border:1px solid #d0d7de;border-radius:6px;
                margin-bottom:16px;">
        <div style="width:8px;height:8px;border-radius:50%;background:{status_color};
                    box-shadow:0 0 6px {status_color};"></div>
        <span style="font-size:13px;color:#636c76;">{status_text}</span>
    </div>
    """, unsafe_allow_html=True)

    # Nav: Home
    if st.button("＋  New Repository", use_container_width=True):
        st.session_state.view = "home"
        st.rerun()

    st.markdown('<div class="section-title" style="margin-top:20px;">YOUR REPOSITORIES</div>',
                unsafe_allow_html=True)

    # List known repos
    known = get_known_repos()
    if not known:
        st.markdown('<div style="font-size:13px;color:#636c76;padding:8px 0;">No repositories yet.<br>Add one above.</div>',
                    unsafe_allow_html=True)
    else:
        for repo_info in known:
            repo_key = f"{repo_info['owner']}/{repo_info['repo']}"
            is_active = st.session_state.active_repo == repo_key
            btn_style = "background:#f6f8fa;border:1px solid #0969da;" if is_active else ""
            if st.button(f"📁  {repo_key}", key=f"repo_{repo_key}", use_container_width=True):
                st.session_state.active_repo  = repo_key
                st.session_state.active_branch = repo_info.get("branch", "main")
                st.session_state.view          = "dashboard"
                st.session_state.left_tab      = "design"
                st.rerun()

    st.markdown("---")

    # Recent jobs
    st.markdown('<div class="section-title">RECENT JOBS</div>', unsafe_allow_html=True)
    jobs = get_jobs(5)
    if not jobs:
        st.markdown('<div style="font-size:12px;color:#636c76;">No jobs yet.</div>',
                    unsafe_allow_html=True)
    for job in jobs[:5]:
        status = job.get("status", "unknown")
        icon = {"success": "✅", "running": "🔄", "failed": "❌", "queued": "⏳"}.get(status, "•")
        sha  = job.get("commit_sha", "")[:10]
        st.markdown(f"""
        <div style="padding:6px 0;border-bottom:1px solid #eaeef2;">
            <div style="font-size:12px;color:#636c76;">{icon} <span style="font-family:'JetBrains Mono',monospace;color:#0969da;">{sha}</span></div>
            <div style="font-size:11px;color:#57606a;margin-top:2px;">{status}</div>
        </div>
        """, unsafe_allow_html=True)

    if st.button("🔄 Refresh", use_container_width=True):
        st.rerun()


# ── MAIN CONTENT ──────────────────────────────────────────────────────────────

# ════════════════════════════════════════════════════════════
# VIEW: HOME — URL Input
# ════════════════════════════════════════════════════════════
if st.session_state.view == "home":
    st.markdown("""
    <div style="max-width:700px;margin:60px auto 0 auto;text-align:center;">
        <div class="hero-title">Which repo would you like to document?</div>
        <div class="hero-sub">Paste a GitHub URL to auto-generate a full system design document</div>
    </div>
    """, unsafe_allow_html=True)

    col_l, col_c, col_r = st.columns([1, 2, 1])
    with col_c:
        github_url = st.text_input(
            "",
            placeholder="https://github.com/owner/repo",
            label_visibility="collapsed",
        )
        branch = st.text_input("", value="main", placeholder="Branch (default: main)",
                               label_visibility="collapsed")

        generate_clicked = st.button("🚀  Generate Design Document", use_container_width=True, type="primary")

        if generate_clicked:
            if not github_url or "github.com" not in github_url:
                st.error("Please enter a valid GitHub URL")
            elif not is_running:
                st.error("FastAPI server is not running. Start it with: `uvicorn main:app --port 8001`")
            else:
                with st.spinner("Submitting..."):
                    try:
                        resp = requests.post(
                            f"{API_BASE}/generate",
                            json={"github_url": github_url, "branch": branch or "main"},
                            timeout=10,
                        )
                        if resp.status_code == 200:
                            data = resp.json()
                            repo_key = f"{data['owner']}/{data['repo']}"
                            st.session_state.active_repo   = repo_key
                            st.session_state.active_branch = data.get("branch", "main")
                            st.session_state.active_job_id = data.get("job_id")
                            st.session_state.view          = "generating"
                            st.rerun()
                        else:
                            st.error(f"Error: {resp.text}")
                    except Exception as e:
                        st.error(f"Could not reach server: {e}")

    # Show existing repos as cards
    known = get_known_repos()
    if known:
        st.markdown("---")
        st.markdown('<div style="text-align:center;font-size:13px;color:#636c76;margin-bottom:16px;">OR OPEN AN EXISTING REPOSITORY</div>',
                    unsafe_allow_html=True)
        cols = st.columns(min(3, len(known)))
        for i, repo_info in enumerate(known[:3]):
            repo_key = f"{repo_info['owner']}/{repo_info['repo']}"
            with cols[i % 3]:
                st.markdown(f"""
                <div class="doc-card" style="cursor:pointer;">
                    <div style="font-size:14px;font-weight:600;color:#1f2328;margin-bottom:4px;">
                        📁 {repo_key}
                    </div>
                    <div style="font-size:12px;color:#636c76;">Branch: {repo_info.get('branch','main')}</div>
                    <div style="font-size:11px;color:#57606a;margin-top:4px;">
                        {repo_info['dt'].strftime('%b %d, %Y') if repo_info.get('dt') else 'Unknown date'}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                if st.button(f"Open", key=f"open_{repo_key}", use_container_width=True):
                    st.session_state.active_repo  = repo_key
                    st.session_state.active_branch = repo_info.get("branch", "main")
                    st.session_state.view          = "dashboard"
                    st.session_state.left_tab      = "design"
                    st.rerun()


# ════════════════════════════════════════════════════════════
# VIEW: GENERATING — Loading screen
# ════════════════════════════════════════════════════════════
elif st.session_state.view == "generating":
    repo_key = st.session_state.active_repo or "repository"
    job_id   = st.session_state.active_job_id

    st.markdown(f"""
    <div style="max-width:560px;margin:80px auto 0 auto;text-align:center;">
        <div style="font-size:48px;margin-bottom:16px;">⚙️</div>
        <div class="hero-title" style="font-size:24px;">Generating documentation</div>
        <div class="hero-sub" style="font-size:14px;margin-bottom:32px;">
            Scanning <span style="color:#0969da;font-family:'JetBrains Mono',monospace;">{repo_key}</span>
            — this takes 2-5 minutes
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Show live job status
    if job_id:
        try:
            jr = requests.get(f"{API_BASE}/jobs/{job_id}", timeout=3)
            if jr.status_code == 200:
                job = jr.json()
                status = job.get("status", "queued")
                steps_done = job.get("steps_completed", [])
                current = job.get("current_step", "")

                all_steps = [
                    ("fetch_repo_context",        "Fetching repository files"),
                    ("detect_repo_type",           "Detecting repo type"),
                    ("analyze_tech_stack",         "Analyzing tech stack"),
                    ("compute_metrics",            "Computing code metrics"),
                    ("generate_repo_description",  "Generating system overview"),
                    ("generate_architecture",      "Generating architecture section"),
                    ("generate_api_section",       "Generating API documentation"),
                    ("generate_data_flow",         "Generating data flow"),
                    ("save_docx",                  "Saving design document (.docx)"),
                    ("push_markdown",              "Pushing DOCUMENTATION.md to repo"),
                    ("save_section_map",           "Saving section map"),
                ]

                col_l2, col_c2, col_r2 = st.columns([1, 2, 1])
                with col_c2:
                    progress = len(steps_done) / max(len(all_steps), 1)
                    st.progress(progress)
                    for step_key, step_label in all_steps:
                        if step_key in steps_done:
                            st.markdown(f'<div class="loading-step done">✅ {step_label}</div>',
                                        unsafe_allow_html=True)
                        elif step_key == current:
                            st.markdown(f'<div class="loading-step active">🔄 {step_label}...</div>',
                                        unsafe_allow_html=True)
                        else:
                            st.markdown(f'<div class="loading-step">○ {step_label}</div>',
                                        unsafe_allow_html=True)

                    if status == "success":
                        st.success("✅ Documentation generated successfully!")
                        if st.button("📂 View Documentation", type="primary", use_container_width=True):
                            st.session_state.view     = "dashboard"
                            st.session_state.left_tab = "design"
                            st.rerun()
                    elif status == "failed":
                        err = job.get("error", "Unknown error")
                        st.error(f"❌ Generation failed: {err}")
                        if st.button("← Try Again"):
                            st.session_state.view = "home"
                            st.rerun()
        except Exception as e:
            st.info(f"Checking job status... ({e})")
    else:
        col_l2, col_c2, col_r2 = st.columns([1, 2, 1])
        with col_c2:
            st.info("Job submitted. Waiting for updates...")


# ════════════════════════════════════════════════════════════
# VIEW: DASHBOARD — Two-panel layout
# ════════════════════════════════════════════════════════════
elif st.session_state.view == "dashboard":
    repo_key = st.session_state.active_repo or ""
    owner, repo_name = (repo_key.split("/", 1) + [""])[:2]
    branch = st.session_state.active_branch

    docs = get_repo_docs(owner, repo_name)
    initial_docs = docs["initial"]
    commit_docs  = docs["commits"]

    # ── Top bar ───────────────────────────────────────────────────────────────
    col_info, col_actions = st.columns([3, 1])
    with col_info:
        st.markdown(f"""
        <div style="display:flex;align-items:center;gap:12px;padding:8px 0 16px 0;">
            <div class="repo-badge">📁 {repo_key}</div>
            <div style="font-size:12px;color:#636c76;padding:3px 8px;background:#ddf4ff;
                        border:1px solid #d0d7de;border-radius:4px;">
                🌿 {branch}
            </div>
            <div style="font-size:12px;color:#636c76;">
                {len(initial_docs)} design doc{"s" if len(initial_docs)!=1 else ""} · {len(commit_docs)} commit doc{"s" if len(commit_docs)!=1 else ""}
            </div>
        </div>
        """, unsafe_allow_html=True)
    with col_actions:
        github_url = f"https://github.com/{repo_key}/blob/{branch}/DOCUMENTATION.md"
        st.markdown(f"""
        <div style="text-align:right;padding-top:8px;">
            <a href="{github_url}" target="_blank"
               style="display:inline-flex;align-items:center;gap:6px;
                      background:#ffffff;border:1px solid #d0d7de;border-radius:6px;
                      padding:6px 14px;text-decoration:none;color:#1f2328;font-size:13px;
                      font-weight:500;transition:border-color 0.2s;">
                🔗 View in GitHub
            </a>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<hr style="margin:0 0 16px 0;">', unsafe_allow_html=True)

    # ── Two-column layout: LEFT nav | RIGHT content ───────────────────────────
    left_col, right_col = st.columns([1, 3])

    with left_col:
        # Section nav
        st.markdown('<div class="section-title">SECTIONS</div>', unsafe_allow_html=True)

        design_btn = st.button("📋  Design Documentation",
                               use_container_width=True,
                               type="primary" if st.session_state.left_tab == "design" else "secondary")
        commits_btn = st.button("🔀  Commit Changes & Impact",
                                use_container_width=True,
                                type="primary" if st.session_state.left_tab == "commits" else "secondary")

        if design_btn:
            st.session_state.left_tab = "design"
            st.rerun()
        if commits_btn:
            st.session_state.left_tab = "commits"
            st.rerun()

        st.markdown("---")

        # Design doc sub-sections
        if st.session_state.left_tab == "design":
            st.markdown('<div class="section-title">ON THIS PAGE</div>', unsafe_allow_html=True)
            sections = [
                "Project Overview",
                "Architecture",
                "Tools & Tech Stack",
                "API Information",
                "Code Quality Metrics",
                "Data Flow",
            ]
            for s in sections:
                st.markdown(f'<div style="font-size:13px;color:#636c76;padding:4px 8px;'
                            f'cursor:pointer;">• {s}</div>', unsafe_allow_html=True)

        # Commit sub-sections
        if st.session_state.left_tab == "commits":
            st.markdown('<div class="section-title">ON THIS PAGE</div>', unsafe_allow_html=True)
            for s in ["Latest Commits", "Impact Analysis", "Files Changed"]:
                st.markdown(f'<div style="font-size:13px;color:#636c76;padding:4px 8px;">• {s}</div>',
                            unsafe_allow_html=True)

    # ── RIGHT PANEL ───────────────────────────────────────────────────────────
    with right_col:

        # ── SECTION 1: DESIGN DOCUMENTATION ──────────────────────────────────
        if st.session_state.left_tab == "design":
            st.markdown("""
            <div style="font-size:24px;font-weight:700;color:#1f2328;
                        letter-spacing:-0.02em;margin-bottom:4px;">
                📋 Design Documentation
            </div>
            <div style="font-size:14px;color:#636c76;margin-bottom:24px;">
                Auto-generated system design document for this repository
            </div>
            """, unsafe_allow_html=True)

            if not initial_docs:
                st.info("No design document generated yet. The document will appear here after generation completes.")
            else:
                latest = initial_docs[0]
                filepath = os.path.join(DOCS_FOLDER, latest["file"])

                # Doc metadata banner
                st.markdown(f"""
                <div style="background:#ddf4ff;border:1px solid #0969da;border-radius:8px;
                            padding:16px 20px;margin-bottom:20px;display:flex;
                            align-items:center;justify-content:space-between;">
                    <div>
                        <div style="font-size:13px;color:#0969da;font-weight:500;margin-bottom:2px;">
                            Latest Design Document
                        </div>
                        <div style="font-size:12px;color:#636c76;">
                            Generated: {latest['dt'].strftime('%B %d, %Y at %H:%M') if latest.get('dt') else 'Unknown'}
                            &nbsp;·&nbsp; Branch: {latest.get('branch','main')}
                        </div>
                    </div>
                    <div style="font-size:11px;color:#1a7f37;background:#dafbe1;
                                border:1px solid #1a7f37;border-radius:12px;padding:3px 10px;">
                        ✓ Up to date
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # ── Parse docx directly into sections ────────────────────────
                SECTION_ORDER = [
                    "Project Overview",
                    "Architecture",
                    "Tools & Tech Stack",
                    "API Information",
                    "Code Quality Metrics",
                    "Data Flow",
                ]
                try:
                    sections_map = read_docx_sections(filepath)
                except Exception as _e:
                    st.error(f"Could not parse document: {_e}")
                    sections_map = {s: [] for s in SECTION_ORDER}


                # ── Display each section ──────────────────────────────────────
                icons = {
                    "Project Overview": "🏠",
                    "Architecture":     "🏗️",
                    "Tools & Tech Stack": "🛠️",
                    "API Information":  "🔌",
                    "Code Quality Metrics": "📊",
                    "Data Flow":        "🔄",
                }

                def render_section_content(sec_lines: list):
                    """Render section lines: sub-headings, bullets, numbered lists, code, tables."""
                    pending_text  = []
                    pending_table = []
                    pending_code  = []
                    in_table = False
                    in_code  = False

                    def flush_text():
                        if pending_text:
                            block = "\n".join(pending_text).strip()
                            if block:
                                st.markdown(block)
                            pending_text.clear()

                    def flush_code():
                        if pending_code:
                            code_block = "\n".join(pending_code)
                            st.code(code_block, language="")
                            pending_code.clear()

                    def flush_table():
                        if not pending_table:
                            return
                        rows = []
                        for tl in pending_table:
                            tl_s = tl.strip()
                            if re.match(r"^\|[-| :]+\|$", tl_s):
                                continue
                            if tl_s.startswith("|") and tl_s.endswith("|"):
                                parts = [p.strip() for p in tl_s.strip("|").split("|")]
                                if parts and any(p for p in parts):
                                    rows.append(parts)
                        if len(rows) > 1:
                            try:
                                import pandas as pd
                                df = pd.DataFrame(rows[1:], columns=rows[0])
                                st.dataframe(df, use_container_width=True, hide_index=True)
                            except Exception:
                                st.markdown("\n".join(pending_table))
                        elif rows:
                            st.markdown("\n".join(pending_table))
                        pending_table.clear()

                    for line in sec_lines:
                        stripped = line.strip()

                        # Inline code line from mermaid block (backtick-wrapped)
                        is_code_line = stripped.startswith("`") and stripped.endswith("`") and len(stripped) > 2

                        # Sub-heading line — bold label on its own
                        is_subheading = (
                            stripped.startswith("**") and stripped.endswith("**")
                            and "\n" not in stripped
                            and not stripped.startswith("**`")
                            and len(stripped) < 100
                            and not stripped[2:-2].strip().startswith("-")
                        )

                        # Real markdown table line
                        is_table_line = (
                            stripped.startswith("|") and stripped.endswith("|")
                            and stripped.count("|") >= 3
                        )

                        if is_code_line:
                            # Flush text/table, accumulate code lines
                            if not in_code:
                                flush_text()
                                flush_table()
                                in_table = False
                                in_code = True
                            pending_code.append(stripped[1:-1])  # strip backticks
                            continue

                        if in_code and not is_code_line:
                            flush_code()
                            in_code = False

                        if is_subheading:
                            flush_text()
                            flush_table()
                            in_table = False
                            # Render sub-heading as its own styled element
                            label = stripped[2:-2].strip()
                            st.markdown(
                                f'<div style="font-size:13px;font-weight:600;'
                                f'color:var(--color-text-primary,#1f2328);'
                                f'margin:14px 0 6px 0;padding-bottom:4px;'
                                f'border-bottom:1px solid #d0d7de;">{label}</div>',
                                unsafe_allow_html=True
                            )
                        elif is_table_line:
                            if not in_table:
                                flush_text()
                                in_table = True
                            pending_table.append(line)
                        else:
                            if in_table:
                                flush_table()
                                in_table = False
                            # Skip blank lines between sub-heading and first bullet
                            if stripped or pending_text:
                                pending_text.append(line)

                    flush_text()
                    flush_table()
                    flush_code()

                for sec_name in SECTION_ORDER:
                    sec_lines = sections_map[sec_name]
                    if not any(l.strip() for l in sec_lines):
                        continue
                    with st.expander(
                        f"{icons.get(sec_name, '•')} {sec_name}",
                        expanded=(sec_name == "Project Overview")
                    ):
                        render_section_content(sec_lines)

                # Download button
                st.markdown("---")
                col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])
                with col_dl1:
                    if os.path.exists(filepath):
                        with open(filepath, "rb") as fh:
                            st.download_button(
                                label="⬇️  Download Design Doc (.docx)",
                                data=fh.read(),
                                file_name=latest["file"],
                                key=f"dl_design_{latest['file']}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )
                    else:
                        st.warning("Document file not found.")
                with col_dl2:
                    st.markdown(f"""
                    <a href="https://github.com/{repo_key}/blob/{branch}/DOCUMENTATION.md"
                       target="_blank"
                       style="display:block;text-align:center;background:#ffffff;
                              border:1px solid #d0d7de;border-radius:6px;padding:8px;
                              text-decoration:none;color:#1f2328;font-size:14px;">
                        🔗 View in GitHub Repo
                    </a>
                    """, unsafe_allow_html=True)

                # Older versions
                if len(initial_docs) > 1:
                    with st.expander(f"📁 Older versions ({len(initial_docs)-1})"):
                        for old_doc in initial_docs[1:]:
                            old_path = os.path.join(DOCS_FOLDER, old_doc["file"])
                            c1, c2 = st.columns([3, 1])
                            with c1:
                                dt_str = old_doc['dt'].strftime('%b %d, %Y %H:%M') if old_doc.get('dt') else 'Unknown'
                                st.markdown(f'<div style="font-size:13px;color:#636c76;">{dt_str}</div>',
                                            unsafe_allow_html=True)
                            with c2:
                                if os.path.exists(old_path):
                                    with open(old_path, "rb") as fh:
                                        st.download_button("⬇", data=fh.read(),
                                                           file_name=old_doc["file"],
                                                           key=f"old_{old_doc['file']}",
                                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # ── SECTION 2: COMMIT CHANGES & IMPACT ANALYSIS ───────────────────────
        elif st.session_state.left_tab == "commits":
            st.markdown("""
            <div style="font-size:24px;font-weight:700;color:#1f2328;
                        letter-spacing:-0.02em;margin-bottom:4px;">
                🔀 Commit Changes & Impact Analysis
            </div>
            <div style="font-size:14px;color:#636c76;margin-bottom:24px;">
                Auto-triggered on every GitHub push — shows what changed and its impact
            </div>
            """, unsafe_allow_html=True)

            if not commit_docs:
                st.markdown("""
                <div class="content-section" style="text-align:center;padding:48px;">
                    <div style="font-size:32px;margin-bottom:12px;">🔗</div>
                    <div style="font-size:16px;color:#1f2328;font-weight:500;margin-bottom:8px;">
                        No commit updates yet
                    </div>
                    <div style="font-size:14px;color:#636c76;">
                        Set up a GitHub webhook pointing to<br>
                        <span style="font-family:'JetBrains Mono',monospace;color:#0969da;">
                            POST http://your-server:8001/webhook
                        </span><br><br>
                        Every push will automatically generate a commit analysis report here.
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                # Show each commit doc
                for i, commit_doc in enumerate(commit_docs[:10]):
                    sha = commit_doc.get("sha", "unknown")
                    branch_c = commit_doc.get("branch", "main")
                    dt_str = commit_doc["dt"].strftime("%b %d, %Y at %H:%M") if commit_doc.get("dt") else "Unknown"
                    filepath = os.path.join(DOCS_FOLDER, commit_doc["file"])

                    with st.expander(
                        f"🔀 Commit `{sha[:7]}` — {dt_str}",
                        expanded=(i == 0)
                    ):
                        # Commit header
                        st.markdown(f"""
                        <div style="display:flex;align-items:center;gap:12px;
                                    padding:12px;background:#f6f8fa;border-radius:6px;
                                    margin-bottom:16px;">
                            <div class="sha-badge">{sha[:7]}</div>
                            <div style="font-size:12px;color:#636c76;">
                                🌿 {branch_c} &nbsp;·&nbsp; 📅 {dt_str}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        if os.path.exists(filepath):
                            doc_text = read_docx_text(filepath)

                            # Parse tables from doc
                            lines = doc_text.split("\n")

                            # Extract sections
                            sections_updated = []
                            commit_table_lines = []
                            impact_table_lines = []
                            in_commit = False
                            in_impact = False

                            for line in lines:
                                ll = line.lower()
                                if "sections updated" in ll:
                                    in_commit = False; in_impact = False
                                elif "commit change" in ll:
                                    in_commit = True; in_impact = False
                                elif "impact analysis" in ll:
                                    in_commit = False; in_impact = True
                                elif in_commit and "|" in line:
                                    commit_table_lines.append(line)
                                elif in_impact and "|" in line:
                                    impact_table_lines.append(line)
                                elif "updated:" in ll or "unchanged:" in ll:
                                    sections_updated.append(line.strip())

                            # Sections updated badges
                            if sections_updated:
                                st.markdown("**Sections Updated**")
                                badge_html = ""
                                for s in sections_updated:
                                    color = "#1a4731" if "updated" in s.lower() else "#1c1c1c"
                                    text_color = "#1a7f37" if "updated" in s.lower() else "#57606a"
                                    badge_html += f'<span style="background:{color};color:{text_color};border-radius:4px;padding:2px 8px;font-size:12px;margin-right:6px;">{s}</span>'
                                st.markdown(badge_html, unsafe_allow_html=True)
                                st.markdown("")

                            # Commit change details table
                            if commit_table_lines:
                                st.markdown("**📝 Files Changed**")
                                import pandas as pd
                                rows = []
                                for row_line in commit_table_lines:
                                    if "|" in row_line and not re.match(r"^[-| :]+$", row_line):
                                        parts = [p.strip() for p in row_line.strip("|").split("|")]
                                        if parts and any(p for p in parts):
                                            rows.append(parts)
                                if len(rows) > 1:
                                    try:
                                        df = pd.DataFrame(rows[1:], columns=rows[0])
                                        st.dataframe(df, use_container_width=True, hide_index=True)
                                    except:
                                        st.code("\n".join(commit_table_lines))
                                st.markdown("")

                            # Impact analysis table
                            if impact_table_lines:
                                st.markdown("**⚡ Impact Analysis**")
                                rows = []
                                for row_line in impact_table_lines:
                                    if "|" in row_line and not re.match(r"^[-| :]+$", row_line):
                                        parts = [p.strip() for p in row_line.strip("|").split("|")]
                                        if parts and any(p for p in parts):
                                            rows.append(parts)
                                if len(rows) > 1:
                                    try:
                                        import pandas as pd
                                        df = pd.DataFrame(rows[1:], columns=rows[0])
                                        st.dataframe(df, use_container_width=True, hide_index=True)
                                    except:
                                        st.code("\n".join(impact_table_lines))
                            else:
                                # Show raw text if no tables found
                                preview = "\n".join(lines[:30])
                                if preview.strip():
                                    st.text(preview)

                            # Download this commit doc
                            st.markdown("---")
                            with open(filepath, "rb") as fh:
                                st.download_button(
                                    label=f"⬇️ Download Commit Report — {sha[:7]}",
                                    data=fh.read(),
                                    file_name=commit_doc["file"],
                                    key=f"dl_commit_{sha}_{i}",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )
                        else:
                            st.warning("Document file not found on disk.")

                # Webhook setup reminder
                st.markdown("""
                <div style="background:#ffffff;border:1px solid #d0d7de;border-radius:8px;
                            padding:16px;margin-top:16px;">
                    <div style="font-size:13px;font-weight:600;color:#1f2328;margin-bottom:8px;">
                        🔗 Webhook Setup
                    </div>
                    <div style="font-size:12px;color:#636c76;">
                        New commits are detected automatically via GitHub webhook.<br>
                        Payload URL: <span style="font-family:'JetBrains Mono',monospace;color:#0969da;">
                        http://your-server:8001/webhook</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)