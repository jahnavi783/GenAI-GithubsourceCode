"""
DOCX builder — shared between initial pipeline and webhook pipeline.

build_initial_docx(ctx)     → full design document from scratch
build_incremental_docx(ctx) → commit-level update doc with changelog
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ── Colour palette ────────────────────────────────────────────────────────────
CLR_DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
CLR_MID_BLUE  = RGBColor(0x2E, 0x74, 0xB5)
CLR_ACCENT    = RGBColor(0x00, 0x70, 0xC0)
CLR_TBL_HDR   = "1F4E79"
CLR_TBL_ALT   = "DEEAF1"
CLR_WHITE     = "FFFFFF"
CLR_BORDER    = "BFBFBF"


# =============================================================================
# PUBLIC API
# =============================================================================

def build_initial_docx(ctx: dict) -> str:
    """
    Builds a full system design document.
    Returns the saved file path.
    """
    doc = Document()
    _set_default_font(doc)

    repo = ctx.get("repo", os.getenv("GITHUB_REPO", "Unknown"))
    owner = ctx.get("owner", os.getenv("GITHUB_OWNER", ""))
    branch = ctx.get("branch", "main")
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # ── Cover ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"System Design Document")
    run.bold = True; run.font.size = Pt(24); run.font.color.rgb = CLR_DARK_BLUE; run.font.name = "Arial"

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"{owner}/{repo}")
    sub_run.font.size = Pt(14); sub_run.font.color.rgb = CLR_MID_BLUE; sub_run.font.name = "Arial"

    doc.add_paragraph()

    # ── Document metadata table ───────────────────────────────────────────────
    _add_heading(doc, "Document Information", level=1)
    meta = [
        ("Repository",  f"{owner}/{repo}"),
        ("Branch",      branch),
        ("Generated",   now_str),
        ("Repo Type",   ctx.get("repo_type", "N/A").upper()),
    ]
    tbl = _make_table(doc, rows=len(meta), cols=2, col_widths=[2160, 7200])
    for i, (label, value) in enumerate(meta):
        fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
        _cell(tbl, i, 0, label, bold=True, fill=fill)
        _cell(tbl, i, 1, value, fill=fill)
    doc.add_paragraph()

    # ── Overview / Description ────────────────────────────────────────────────
    repo_desc = ctx.get("repo_description", "")
    if repo_desc:
        _render_llm_bullet_sections(doc, repo_desc)

    # ── Architecture ──────────────────────────────────────────────────────────
    architecture = ctx.get("architecture", "")
    if architecture:
        _add_heading(doc, "Architecture", level=1)
        _render_llm_bullet_sections(doc, architecture)

    # ── Tech Stack ────────────────────────────────────────────────────────────
    tech = ctx.get("tech_stack", {})
    if tech:
        _add_heading(doc, "Tools & Tech Stack", level=1)
        langs = tech.get("languages", [])
        if langs:
            _add_heading(doc, "Languages", level=2)
            for lang in langs:
                _bullet(doc, lang)
        libs = tech.get("frameworks_and_libs", [])
        if libs:
            _add_heading(doc, "Frameworks & Libraries", level=2)
            lib_tbl = _make_table(doc, rows=1 + len(libs), cols=2, col_widths=[4680, 4680])
            _cell(lib_tbl, 0, 0, "Library / Framework", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE)
            _cell(lib_tbl, 0, 1, "Category",            bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE)
            for i, lib in enumerate(libs, start=1):
                fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
                _cell(lib_tbl, i, 0, lib.get("name", ""),     fill=fill)
                _cell(lib_tbl, i, 1, lib.get("category", ""), fill=fill)
        infra = tech.get("infra", [])
        if infra:
            _add_heading(doc, "Infrastructure", level=2)
            for item in infra:
                _bullet(doc, item)
        doc.add_paragraph()

    # ── Code Quality Metrics ──────────────────────────────────────────────────
    metrics = ctx.get("metrics", [])
    if metrics:
        _add_heading(doc, "Code Quality Metrics", level=1)
        m_tbl = _make_table(doc, rows=1 + len(metrics), cols=3, col_widths=[4500, 2460, 2400])
        _cell(m_tbl, 0, 0, "Metric", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
        _cell(m_tbl, 0, 1, "Value",  bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
        _cell(m_tbl, 0, 2, "Status", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
        for i, m in enumerate(metrics, start=1):
            fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
            _cell(m_tbl, i, 0, m.get("metric", ""),     fill=fill, font_size=10)
            _cell(m_tbl, i, 1, str(m.get("value", "")), fill=fill, font_size=10)
            _cell(m_tbl, i, 2, m.get("status", ""),     fill=fill, font_size=10)
        doc.add_paragraph()

    # ── API Endpoints ─────────────────────────────────────────────────────────
    api_section = ctx.get("api_section", "")
    if api_section:
        _add_heading(doc, "API Endpoints", level=1)
        _render_llm_bullet_sections(doc, api_section)

    # ── Data Flow ─────────────────────────────────────────────────────────────
    data_flow = ctx.get("data_flow", "")
    if data_flow:
        _add_heading(doc, "Data Flow", level=1)
        _render_llm_bullet_sections(doc, data_flow)

    # ── Save ──────────────────────────────────────────────────────────────────
    return _save_doc(doc, f"design_{owner}_{repo}_{branch}")


def build_incremental_docx(ctx: dict) -> str:
    """
    Builds a commit-level update document showing what changed.
    Returns the saved file path.
    """
    import re

    doc = Document()
    _set_default_font(doc)

    repo = ctx.get("repo", os.getenv("GITHUB_REPO", "Unknown"))
    owner = ctx.get("owner", os.getenv("GITHUB_OWNER", ""))
    branch = ctx.get("branch", "main")
    sha = ctx.get("commit_sha", "")[:7]
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Commit Documentation Report")
    run.bold = True; run.font.size = Pt(22); run.font.color.rgb = CLR_DARK_BLUE; run.font.name = "Arial"
    doc.add_paragraph()

    # ── Commit info table ─────────────────────────────────────────────────────
    _add_heading(doc, "Commit Information", level=1)
    meta = [
        ("Repository",      f"{owner}/{repo}"),
        ("Branch",          branch),
        ("Commit SHA",      ctx.get("commit_sha", "N/A")[:7]),
        ("Author",          ctx.get("author", "N/A")),
        ("Commit Message",  ctx.get("message", "N/A")),
        ("Generated",       now_str),
    ]
    tbl = _make_table(doc, rows=len(meta), cols=2, col_widths=[2160, 7200])
    for i, (label, value) in enumerate(meta):
        fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
        _cell(tbl, i, 0, label, bold=True, fill=fill)
        _cell(tbl, i, 1, value, fill=fill)
    doc.add_paragraph()

    # ── Sections updated in this commit ──────────────────────────────────────
    affected = ctx.get("affected_sections", [])
    if affected:
        _add_heading(doc, "Sections Updated This Commit", level=1)
        unchanged = [s for s in ["repo_description", "architecture", "api_section", "data_flow"]
                     if s not in affected]
        for s in affected:
            _bullet(doc, f"✅  Updated: {s.replace('_', ' ').title()}")
        for s in unchanged:
            _bullet(doc, f"—  Unchanged: {s.replace('_', ' ').title()}")
        doc.add_paragraph()

    # ── Updated sections content ──────────────────────────────────────────────
    if "repo_description" in affected and ctx.get("repo_description"):
        _render_llm_bullet_sections(doc, ctx["repo_description"])

    if "architecture" in affected and ctx.get("architecture"):
        _add_heading(doc, "Architecture (Updated)", level=1)
        _render_llm_bullet_sections(doc, ctx["architecture"])

    if "api_section" in affected and ctx.get("api_section"):
        _add_heading(doc, "API Endpoints (Updated)", level=1)
        _render_llm_bullet_sections(doc, ctx["api_section"])

    if "data_flow" in affected and ctx.get("data_flow"):
        _add_heading(doc, "Data Flow (Updated)", level=1)
        _render_llm_bullet_sections(doc, ctx["data_flow"])

    # ── Commit change details ─────────────────────────────────────────────────
    doc_raw = ctx.get("documentation", "")
    if doc_raw:
        _add_heading(doc, "Commit Change Details", level=1)
        commit_rows = _parse_markdown_table(doc_raw)
        COMMIT_HEADERS = ["File Changed", "Change Type", "Description", "Lines Added", "Lines Removed", "Risk Level"]
        if commit_rows:
            data_rows = commit_rows[1:] if len(commit_rows) > 1 else commit_rows
            widths = _distribute_widths(len(COMMIT_HEADERS), 9360)
            c_tbl = _make_table(doc, rows=1 + len(data_rows), cols=len(COMMIT_HEADERS), col_widths=widths)
            for ci, h in enumerate(COMMIT_HEADERS):
                _cell(c_tbl, 0, ci, h, bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=8)
            for ri, row in enumerate(data_rows, start=1):
                fill = CLR_TBL_ALT if ri % 2 == 0 else CLR_WHITE
                for ci in range(len(COMMIT_HEADERS)):
                    val = row[ci] if ci < len(row) else ""
                    _cell(c_tbl, ri, ci, val, fill=fill, font_size=8)
        doc.add_paragraph()

    # ── Impact analysis ───────────────────────────────────────────────────────
    impact_raw = ctx.get("impact_analysis", "")
    if impact_raw:
        _add_heading(doc, "Impact Analysis", level=1)
        impact_rows = _parse_markdown_table(impact_raw)
        IMPACT_HEADERS = ["Area Impacted", "Type of Impact", "Severity", "Description", "Action Required"]
        if impact_rows:
            data_rows = impact_rows[1:] if len(impact_rows) > 1 else impact_rows
            widths = _distribute_widths(len(IMPACT_HEADERS), 9360)
            i_tbl = _make_table(doc, rows=1 + len(data_rows), cols=len(IMPACT_HEADERS), col_widths=widths)
            for ci, h in enumerate(IMPACT_HEADERS):
                _cell(i_tbl, 0, ci, h, bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
            for ri, row in enumerate(data_rows, start=1):
                fill = CLR_TBL_ALT if ri % 2 == 0 else CLR_WHITE
                for ci in range(len(IMPACT_HEADERS)):
                    val = row[ci] if ci < len(row) else ""
                    _cell(i_tbl, ri, ci, val, fill=fill, font_size=10)
        doc.add_paragraph()

    # ── Tech stack (always shown) ─────────────────────────────────────────────
    tech = ctx.get("tech_stack", {})
    if tech:
        _add_heading(doc, "Tools & Tech Stack", level=1)
        langs = tech.get("languages", [])
        if langs:
            _add_heading(doc, "Languages", level=2)
            for lang in langs:
                _bullet(doc, lang)
        libs = tech.get("frameworks_and_libs", [])
        if libs:
            _add_heading(doc, "Frameworks & Libraries", level=2)
            lib_tbl = _make_table(doc, rows=1 + len(libs), cols=2, col_widths=[4680, 4680])
            _cell(lib_tbl, 0, 0, "Library / Framework", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE)
            _cell(lib_tbl, 0, 1, "Category",            bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE)
            for i, lib in enumerate(libs, start=1):
                fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
                _cell(lib_tbl, i, 0, lib.get("name", ""),     fill=fill)
                _cell(lib_tbl, i, 1, lib.get("category", ""), fill=fill)
        doc.add_paragraph()

    # ── Metrics ───────────────────────────────────────────────────────────────
    metrics = ctx.get("metrics", [])
    if metrics:
        _add_heading(doc, "Code Quality Metrics", level=1)
        m_tbl = _make_table(doc, rows=1 + len(metrics), cols=3, col_widths=[4500, 2460, 2400])
        _cell(m_tbl, 0, 0, "Metric", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
        _cell(m_tbl, 0, 1, "Value",  bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
        _cell(m_tbl, 0, 2, "Status", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
        for i, m in enumerate(metrics, start=1):
            fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
            _cell(m_tbl, i, 0, m.get("metric", ""),     fill=fill, font_size=10)
            _cell(m_tbl, i, 1, str(m.get("value", "")), fill=fill, font_size=10)
            _cell(m_tbl, i, 2, m.get("status", ""),     fill=fill, font_size=10)
        doc.add_paragraph()

    return _save_doc(doc, f"commit_{sha}_{branch}")


# =============================================================================
# INTERNAL HELPERS
# =============================================================================

def _save_doc(doc: Document, prefix: str) -> str:
    docs_folder = os.getenv("DOCS_FOLDER", "./generated_docs")
    os.makedirs(docs_folder, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{prefix}_{ts}.docx"
    filepath = os.path.join(docs_folder, filename)
    doc.save(filepath)
    return filepath


def _set_default_font(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)


def _add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    if level == 1:
        run.font.size = Pt(14); run.font.color.rgb = CLR_DARK_BLUE
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(12); run.font.color.rgb = CLR_MID_BLUE
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(2)
    else:
        run.font.size = Pt(11); run.font.color.rgb = CLR_ACCENT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(2)


def _bullet(doc: Document, text: str):
    import re
    para = doc.add_paragraph(style="List Bullet")
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def _numbered_bullet(doc: Document, number: str, text: str):
    """Render a numbered step as a formatted paragraph with the number badge."""
    import re
    para = doc.add_paragraph(style="List Number")
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def _make_table(doc: Document, rows: int, cols: int, col_widths: list) -> object:
    from lxml import etree
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    total_width = sum(col_widths)
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    for old_w in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old_w)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    existing_grid = tbl_el.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl_el.remove(existing_grid)
    tbl_el.append(tbl_grid)
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(col_widths[j]))
            tc_w.set(qn("w:type"), "dxa")
            existing_tcw = tc_pr.find(qn("w:tcW"))
            if existing_tcw is not None:
                tc_pr.remove(existing_tcw)
            tc_pr.insert(0, tc_w)
    return tbl


def _cell(tbl, row: int, col: int, text: str, bold: bool = False,
          fill: str = None, font_color=None, font_size: float = 10):
    cell = tbl.cell(row, col)
    if fill:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    parts = str(text).split("<br>")
    for pi, part in enumerate(parts):
        if pi > 0:
            run = para.add_run()
            run.add_break()
        run = para.add_run(part.strip())
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        if font_color:
            if isinstance(font_color, str):
                r = int(font_color[0:2], 16)
                g = int(font_color[2:4], 16)
                b = int(font_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            else:
                run.font.color.rgb = font_color


def _distribute_widths(cols: int, total: int = 9360) -> list:
    base = total // cols
    remainder = total % cols
    return [base + (1 if i < remainder else 0) for i in range(cols)]


def _parse_markdown_table(text: str) -> list[list[str]]:
    import re
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-| :]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


def _render_llm_bullet_sections(doc: Document, text: str):
    import re
    in_code = False
    is_mermaid = False
    mermaid_lines: list[str] = []
    current_heading = None
    SKIP_SECTIONS = {"src folder", "environment configuration (.env)", "environment configuration"}
    lines = text.splitlines()
    i = 0
    skip_section = False

    while i < len(lines):
        stripped = lines[i].strip()
        if not in_code and stripped.startswith("```"):
            in_code = True
            is_mermaid = "mermaid" in stripped.lower()
            mermaid_lines = []
            i += 1
            continue
        if in_code and stripped.startswith("```"):
            in_code = False
            if mermaid_lines:
                _render_mermaid_block(doc, mermaid_lines)
            is_mermaid = False
            mermaid_lines = []
            i += 1
            continue
        if in_code:
            mermaid_lines.append(lines[i])
            i += 1
            continue
        if stripped.startswith("## "):
            current_heading = stripped[3:].strip()
            skip_section = current_heading.lower() in SKIP_SECTIONS
            if not skip_section:
                _add_heading(doc, current_heading, level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            if not skip_section:
                _add_heading(doc, stripped[4:].strip(), level=2)
            i += 1
            continue
        if skip_section:
            i += 1
            continue
        if stripped.startswith(("- ", "* ", "• ")):
            _bullet(doc, stripped[2:].strip())
            i += 1
            continue
        # Numbered list: "1. ", "2. ", etc.
        num_match = __import__('re').match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            _numbered_bullet(doc, num_match.group(1), num_match.group(2).strip())
            i += 1
            continue
        if stripped and current_heading:
            _bullet(doc, stripped)
            i += 1
            continue
        if not stripped:
            doc.add_paragraph()
        i += 1

    if in_code and mermaid_lines:
        _render_mermaid_block(doc, mermaid_lines)


def _render_mermaid_block(doc: Document, mermaid_lines: list[str]):
    from docx.oxml.ns import qn as _qn
    # Label paragraph
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after  = Pt(0)
    lr = label.add_run("Architecture Diagram")
    lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(9); lr.font.color.rgb = CLR_ACCENT

    # Each mermaid line gets its own paragraph (avoids \n in runs which corrupts Word)
    for idx, line in enumerate(mermaid_lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        pPr = para._p.get_or_add_pPr()
        # Grey background on every line
        shd = OxmlElement("w:shd")
        shd.set(_qn("w:val"), "clear"); shd.set(_qn("w:color"), "auto"); shd.set(_qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        # Blue border only on first and last lines
        pBdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            bdr = OxmlElement(f"w:{side}")
            bdr.set(_qn("w:val"), "single"); bdr.set(_qn("w:sz"), "4")
            bdr.set(_qn("w:space"), "4"); bdr.set(_qn("w:color"), "2E74B5")
            pBdr.append(bdr)
        pPr.append(pBdr)
        run = para.add_run(line)
        run.font.name = "Courier New"; run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()
