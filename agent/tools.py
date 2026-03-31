import os
import env_loader  # noqa: F401 — loads .env with absolute path
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt, Inches

from analyzers.metrics import compute_metrics
from analyzers.repo_type_detector import detect_repo_type
from analyzers.tech_stack_analyzer import analyze_tech_stack
from github.fetcher import fetch_commit_diff, fetch_repo_context
from github.pusher import push_markdown_to_repo
from llm.model import get_llm
from llm.prompts import DOC_TEMPLATE, IMPACT_ANALYSIS_TEMPLATE, REPO_DESCRIPTION_TEMPLATE
from vectordb.store import retrieve_similar, store_commit


# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────────────────────────────────────
CLR_DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)   # title / h1
CLR_MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)   # h2
CLR_ACCENT     = RGBColor(0x00, 0x70, 0xC0)   # h3 / section headers
CLR_TBL_HDR    = "1F4E79"                       # dark blue fill (hex, no #)
CLR_TBL_ALT    = "DEEAF1"                       # light blue alt row
CLR_WHITE      = "FFFFFF"
CLR_BORDER     = "BFBFBF"


def make_steps(ctx: dict):

    # ── STEP 1 ────────────────────────────────────────────────────────────────
    def fetch_diff():
        data = fetch_commit_diff(ctx["commit_sha"])
        ctx.update(data)
        print(f"Fetched diff for {ctx['commit_sha'][:7]}")

    # ── STEP 2 ────────────────────────────────────────────────────────────────
    def fetch_repo_ctx():
        repo_data = fetch_repo_context()
        ctx["file_tree"] = repo_data["file_tree"]
        ctx["key_file_contents"] = repo_data["key_file_contents"]
        print(f"Fetched repo context: {len(ctx['file_tree'])} files")

    # ── STEP 3 ────────────────────────────────────────────────────────────────
    def detect_type():
        ctx["repo_type"] = detect_repo_type(ctx["file_tree"], ctx["key_file_contents"])
        print(f"Repo type: {ctx['repo_type']}")

    # ── STEP 4 ────────────────────────────────────────────────────────────────
    def analyze_stack():
        ctx["tech_stack"] = analyze_tech_stack(ctx["file_tree"], ctx["key_file_contents"])
        print(f"Tech stack: {ctx['tech_stack']['languages']}")

    # ── STEP 5 ────────────────────────────────────────────────────────────────
    def run_metrics():
        ctx["metrics"] = compute_metrics(
            ctx["repo_type"], ctx["file_tree"], ctx["key_file_contents"]
        )
        print(f"Computed {len(ctx['metrics'])} metrics")

    # ── STEP 6 ────────────────────────────────────────────────────────────────
    def retrieve_context():
        ctx["similar_docs"] = retrieve_similar(ctx["diff"], top_k=3)
        print(f"Retrieved {len(ctx['similar_docs'])} similar docs")

    # ── STEP 7 ────────────────────────────────────────────────────────────────
    def generate_repo_description():
        llm = get_llm()
        repo_name = os.getenv("GITHUB_REPO", "Unknown Repo")
        key_files_text = ""
        for filepath, content in ctx["key_file_contents"].items():
            key_files_text += f"\n--- {filepath} ---\n{content[:2000]}\n"
            if len(key_files_text) > 10000:
                break
        # Also include a sample of .tsx/.ts/.py files for real code context
        extra_context = ""
        for filepath in ctx["file_tree"]:
            if filepath not in ctx["key_file_contents"]:
                ext = filepath.rsplit(".", 1)[-1].lower()
                if ext in ("tsx", "ts", "jsx", "js", "py") and "test" not in filepath.lower():
                    try:
                        from github.fetcher import _fetch_file_content
                        import os as _os
                        from github.auth import get_installation_token
                        owner = _os.getenv("GITHUB_OWNER")
                        repo  = _os.getenv("GITHUB_REPO")
                        token = get_installation_token()
                        base  = f"https://api.github.com/repos/{owner}/{repo}"
                        headers = {"Authorization": f"token {token}",
                                   "Accept": "application/vnd.github.v3+json"}
                        content = _fetch_file_content(base, filepath, headers)
                        if content:
                            extra_context += f"\n--- {filepath} ---\n{content[:800]}\n"
                    except Exception:
                        pass
            if len(extra_context) > 4000:
                break
        file_tree_text = "\n".join(ctx["file_tree"][:150])
        prompt = REPO_DESCRIPTION_TEMPLATE.format(
            repo_name=repo_name,
            file_tree=file_tree_text,
            key_files_content=key_files_text + extra_context,
        )
        ctx["repo_description"] = str(llm.invoke(prompt))
        print("Repo description generated")

    # ── STEP 8 ────────────────────────────────────────────────────────────────
    def generate_impact_analysis():
        llm = get_llm()
        repo_overview = ctx.get("repo_description", "")[:2000]
        files_changed = ", ".join(ctx.get("files_changed", []))
        prompt = IMPACT_ANALYSIS_TEMPLATE.format(
            repo_overview=repo_overview,
            files_changed=files_changed,
            diff=ctx.get("diff", "")[:3000],
        )
        ctx["impact_analysis"] = str(llm.invoke(prompt))
        print("Impact analysis generated")

    # ── STEP 9 ────────────────────────────────────────────────────────────────
    def generate_commit_docs():
        llm = get_llm()
        context = "\n---\n".join(
            [item.get("documentation", "") for item in ctx.get("similar_docs", [])]
        )
        prompt = DOC_TEMPLATE.format(
            author=ctx.get("author", "Unknown"),
            branch=ctx.get("branch", "Unknown"),
            sha=ctx.get("commit_sha", "Unknown"),
            timestamp=ctx.get("timestamp", ""),
            files=", ".join(ctx.get("files_changed", [])),
            context=context if context else "No similar past commits found",
            diff=ctx.get("diff", "")[:4000],
        )
        ctx["documentation"] = str(llm.invoke(prompt))
        print("Commit documentation generated")

    # ── STEP 10 ───────────────────────────────────────────────────────────────
    def save_to_vectordb():
        store_commit(
            ctx["commit_sha"], ctx["diff"], ctx["documentation"], ctx["author"]
        )
        print("Saved to ChromaDB")

    # ── STEP 11 ───────────────────────────────────────────────────────────────
    def save_docx():
        doc = Document()
        _set_default_font(doc)

        # ── TITLE ──────────────────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Bridge gaps between design documentation and source code")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = CLR_DARK_BLUE
        run.font.name = "Arial"
        doc.add_paragraph()

        # ── COMMIT INFO TABLE ──────────────────────────────────────────────
        _add_heading(doc, "Commit Information", level=1)
        meta = [
            ("Repository", os.getenv("GITHUB_REPO", "N/A")),
            ("Branch",     ctx.get("branch", "N/A")),
            ("Commit SHA", ctx.get("commit_sha", "N/A")[:7]),
            ("Author",     ctx.get("author", "N/A")),
            ("Generated",  datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        tbl = _make_table(doc, rows=len(meta), cols=2, col_widths=[2160, 7200])
        for i, (label, value) in enumerate(meta):
            _cell(tbl, i, 0, label, bold=True, fill=CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE)
            _cell(tbl, i, 1, value,            fill=CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE)
        doc.add_paragraph()

        # ── REPO DESCRIPTION (bullet sections) ────────────────────────────
        repo_desc = ctx.get("repo_description", "")
        if repo_desc:
            _render_llm_bullet_sections(doc, repo_desc)

        # ── TOOLS & TECH STACK ────────────────────────────────────────────
        tech = ctx.get("tech_stack", {})
        if tech:
            _add_heading(doc, "Tools & Tech Stack", level=1)

            langs = tech.get("languages", [])
            if langs:
                _add_heading(doc, "Languages", level=2)
                for lang in langs:
                    _bullet(doc, lang)

            repo_type = ctx.get("repo_type", "")
            if repo_type:
                _add_heading(doc, "Repository Type", level=2)
                _bullet(doc, repo_type.upper())

            libs = tech.get("frameworks_and_libs", [])
            if libs:
                _add_heading(doc, "Frameworks & Libraries", level=2)
                lib_tbl = _make_table(doc, rows=1 + len(libs), cols=2, col_widths=[4680, 4680])
                _cell(lib_tbl, 0, 0, "Library / Framework", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE)
                _cell(lib_tbl, 0, 1, "Category",            bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE)
                for i, lib in enumerate(libs, start=1):
                    fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
                    _cell(lib_tbl, i, 0, lib["name"],     fill=fill)
                    _cell(lib_tbl, i, 1, lib["category"], fill=fill)

            infra = tech.get("infra", [])
            if infra:
                _add_heading(doc, "Infrastructure", level=2)
                for item in infra:
                    _bullet(doc, item)

            doc.add_paragraph()

        # ── CODE QUALITY METRICS ──────────────────────────────────────────
        metrics = ctx.get("metrics", [])
        if metrics:
            _add_heading(doc, "Code Quality Metrics", level=1)
            # Full page width: Metric=4500, Value=2460, Status=2400  (total=9360)
            m_tbl = _make_table(doc, rows=1 + len(metrics), cols=3,
                                col_widths=[4500, 2460, 2400])
            _cell(m_tbl, 0, 0, "Metric", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
            _cell(m_tbl, 0, 1, "Value",  bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
            _cell(m_tbl, 0, 2, "Status", bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
            for i, m in enumerate(metrics, start=1):
                fill = CLR_TBL_ALT if i % 2 == 0 else CLR_WHITE
                _cell(m_tbl, i, 0, m.get("metric", ""),      fill=fill, font_size=10)
                _cell(m_tbl, i, 1, str(m.get("value", "")),  fill=fill, font_size=10)
                _cell(m_tbl, i, 2, m.get("status", ""),      fill=fill, font_size=10)
            doc.add_paragraph()

        # ── COMBINED COMMIT INFO TABLE ────────────────────────────────────
        doc_raw = ctx.get("documentation", "")
        if doc_raw:
            _add_heading(doc, "Commit Change Details", level=1)
            commit_rows = _parse_markdown_table(doc_raw)
            COMMIT_HEADERS = ["File Changed", "Change Type", "Description", "Lines Added", "Lines Removed", "Risk Level"]
            if commit_rows:
                headers = COMMIT_HEADERS
                data_rows = commit_rows[1:] if len(commit_rows) > 1 else commit_rows
                widths = _distribute_widths(len(headers), 9360)
                c_tbl = _make_table(doc, rows=1 + len(data_rows),
                                    cols=len(headers), col_widths=widths)
                for ci, h in enumerate(headers):
                    _cell(c_tbl, 0, ci, h, bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=8)
                for ri, row in enumerate(data_rows, start=1):
                    fill = CLR_TBL_ALT if ri % 2 == 0 else CLR_WHITE
                    for ci in range(len(headers)):
                        val = row[ci] if ci < len(row) else ""
                        _cell(c_tbl, ri, ci, val, fill=fill, font_size=8)
            doc.add_paragraph()

        # ── IMPACT ANALYSIS TABLE (after commit details) ──────────────────
        impact_raw = ctx.get("impact_analysis", "")
        if impact_raw:
            _add_heading(doc, "Impact Analysis", level=1)
            impact_rows = _parse_markdown_table(impact_raw)
            IMPACT_HEADERS = ["Area Impacted", "Type of Impact", "Severity", "Description", "Action Required"]
            if impact_rows:
                headers = IMPACT_HEADERS
                data_rows = impact_rows[1:] if len(impact_rows) > 1 else impact_rows
                widths = _distribute_widths(len(headers), 9360)
                i_tbl = _make_table(doc, rows=1 + len(data_rows),
                                    cols=len(headers), col_widths=widths)
                for ci, h in enumerate(headers):
                    _cell(i_tbl, 0, ci, h, bold=True, fill=CLR_TBL_HDR, font_color=CLR_WHITE, font_size=10)
                for ri, row in enumerate(data_rows, start=1):
                    fill = CLR_TBL_ALT if ri % 2 == 0 else CLR_WHITE
                    for ci in range(len(headers)):
                        val = row[ci] if ci < len(row) else ""
                        _cell(i_tbl, ri, ci, val, fill=fill, font_size=10)
            doc.add_paragraph()

        # ── SAVE ──────────────────────────────────────────────────────────
        docs_folder = os.getenv("DOCS_FOLDER", "./generated_docs")
        os.makedirs(docs_folder, exist_ok=True)
        filename = (
            f"commit_{ctx['commit_sha'][:7]}_"
            f"{ctx.get('branch', 'unknown')}_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        filepath = os.path.join(docs_folder, filename)
        doc.save(filepath)
        ctx["docx_path"] = filepath
        print(f"DOCX saved: {filepath}")

    # ── STEP 12 ───────────────────────────────────────────────────────────────
    def push_markdown():
        markdown = _build_markdown(ctx)
        url = push_markdown_to_repo(
            markdown, ctx["commit_sha"], ctx.get("branch", "main")
        )
        ctx["markdown_url"] = url
        print(f"Markdown pushed: {url}")

    return {
        "fetch_diff":                fetch_diff,
        "fetch_repo_context":        fetch_repo_ctx,
        "detect_repo_type":          detect_type,
        "analyze_tech_stack":        analyze_stack,
        "compute_metrics":           run_metrics,
        "retrieve_context":          retrieve_context,
        "generate_repo_description": generate_repo_description,
        "generate_impact_analysis":  generate_impact_analysis,
        "generate_commit_docs":      generate_commit_docs,
        "save_vectordb":             save_to_vectordb,
        "save_docx":                 save_docx,
        "push_markdown":             push_markdown,
    }


def make_tools(ctx: dict):
    return make_steps(ctx)


# =============================================================================
# DOCX HELPERS
# =============================================================================

def _set_default_font(doc: Document):
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)


def _add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = CLR_DARK_BLUE
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = CLR_MID_BLUE
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(2)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = CLR_ACCENT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(2)


def _bullet(doc: Document, text: str):
    """Add a bullet point paragraph with bold support for **text** markers."""
    para = doc.add_paragraph(style="List Bullet")
    _add_inline_bold(para, text)


def _add_inline_bold(para, text: str):
    """Parse **bold** markers and add runs accordingly."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(10)


def _make_table(doc: Document, rows: int, cols: int, col_widths: list) -> object:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from lxml import etree

    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"

    total_width = sum(col_widths)
    tbl_el = tbl._tbl  # lxml element, type CT_Tbl

    # ── tblPr: find or create ────────────────────────────────────────────
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)

    # Remove any existing tblW inside tblPr
    for old_w in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old_w)

    # Set total table width
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"),    str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    # ── tblGrid: column widths ───────────────────────────────────────────
    existing_grid = tbl_el.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl_el.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tbl_grid.append(gc)

    # Insert tblGrid right after tblPr
    tbl_pr_idx = list(tbl_el).index(tbl_pr)
    tbl_el.insert(tbl_pr_idx + 1, tbl_grid)

    # ── Per-cell width via XML ────────────────────────────────────────────
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(col_widths):
                continue
            tc = cell._tc

            # find or create tcPr
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)

            # remove old tcW
            for old_w in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(old_w)

            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"),    str(col_widths[ci]))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.insert(0, tc_w)

    return tbl


def _cell(tbl, row: int, col: int, text: str,
          bold: bool = False,
          fill: str = None,
          font_color: object = None,
          font_size: int = 9):
    """Write text into a table cell with optional formatting."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cell = tbl.rows[row].cells[col]

    # Fill / shading
    if fill:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    # Clear existing paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)

    # Handle <br> line breaks
    parts = str(text).split("<br>")
    for pi, part in enumerate(parts):
        if pi > 0:
            run = para.add_run()
            run.add_break()
        run = para.add_run(part.strip())
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        if font_color:
            if isinstance(font_color, str):
                from docx.shared import RGBColor as RC
                r = int(font_color[0:2], 16)
                g = int(font_color[2:4], 16)
                b = int(font_color[4:6], 16)
                run.font.color.rgb = RC(r, g, b)
            else:
                run.font.color.rgb = font_color


def _distribute_widths(cols: int, total: int = 9360) -> list:
    """Distribute total width evenly across columns."""
    base = total // cols
    remainder = total % cols
    return [base + (1 if i < remainder else 0) for i in range(cols)]


# =============================================================================
# LLM OUTPUT PARSERS
# =============================================================================

def _render_llm_bullet_sections(doc: Document, text: str):
    """
    Parse LLM output that uses ## headings and bullet lines.
    Renders headings as Word headings, bullet lines as List Bullet paragraphs,
    and mermaid/code fences as a styled monospaced code block.
    """
    in_code = False
    is_mermaid = False
    mermaid_lines: list[str] = []
    current_heading = None

    # Sections to skip entirely (removed from output)
    SKIP_SECTIONS = {"src folder", "environment configuration (.env)", "environment configuration"}

    lines = text.splitlines()
    i = 0
    skip_section = False

    while i < len(lines):
        stripped = lines[i].strip()

        # ── opening code fence ──────────────────────────────────────────
        if not in_code and stripped.startswith("```"):
            in_code = True
            is_mermaid = "mermaid" in stripped.lower()
            mermaid_lines = []
            i += 1
            continue

        # ── closing code fence ──────────────────────────────────────────
        if in_code and stripped.startswith("```"):
            in_code = False
            if mermaid_lines:
                _render_mermaid_block(doc, mermaid_lines)
            is_mermaid = False
            mermaid_lines = []
            i += 1
            continue

        # ── inside code fence → collect lines ──────────────────────────
        if in_code:
            mermaid_lines.append(lines[i])
            i += 1
            continue

        # ── section heading ─────────────────────────────────────────────
        if stripped.startswith("## "):
            current_heading = stripped[3:].strip()
            # Check if this section should be skipped
            skip_section = current_heading.lower() in SKIP_SECTIONS
            if not skip_section:
                _add_heading(doc, current_heading, level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            if not skip_section:
                _add_heading(doc, stripped[4:].strip(), level=2)
            i += 1
            continue

        # ── skip entire section if flagged ──────────────────────────────
        if skip_section:
            i += 1
            continue

        # ── bullet lines ────────────────────────────────────────────────
        if stripped.startswith(("- ", "* ", "• ")):
            content_text = stripped[2:].strip()
            _bullet(doc, content_text)
            i += 1
            continue

        # ── non-empty plain lines: after mermaid block → style as note ──
        if stripped and current_heading and current_heading.lower() == "codebase structure":
            # Plain sentences after the diagram → render as a styled note paragraph
            if not stripped.startswith(("- ", "* ", "• ")):
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after  = Pt(4)
                # Left border note style
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _qn
                pPr = para._p.get_or_add_pPr()
                pBdr = _OE('w:pBdr')
                left = _OE('w:left')
                left.set(_qn('w:val'), 'single')
                left.set(_qn('w:sz'), '12')
                left.set(_qn('w:space'), '8')
                left.set(_qn('w:color'), 'DEEAF1')
                pBdr.append(left)
                pPr.append(pBdr)
                run = para.add_run(stripped)
                run.font.name  = "Arial"
                run.font.size  = Pt(10)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                i += 1
                continue

        # ── non-empty plain lines in a section → bullet ─────────────────
        if stripped and current_heading:
            _bullet(doc, stripped)
            i += 1
            continue

        # ── blank line ───────────────────────────────────────────────────
        if not stripped:
            doc.add_paragraph()
        i += 1

    # flush unclosed fence
    if in_code and mermaid_lines:
        _render_mermaid_block(doc, mermaid_lines)


def _render_mermaid_block(doc: Document, mermaid_lines: list[str]):
    """
    Renders a mermaid diagram as a styled monospaced code block inside the docx.
    Uses a light grey shaded box with Courier New font so it reads clearly.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Header label
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after  = Pt(0)
    lr = label.add_run("Architecture Diagram")
    lr.bold = True
    lr.font.name = "Arial"
    lr.font.size = Pt(9)
    lr.font.color.rgb = CLR_ACCENT

    # One paragraph per line, all in a shaded box
    content = "\n".join(mermaid_lines)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)

    # Apply grey shading to paragraph
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    run = para.add_run(content)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Border box around the paragraph
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), '2E74B5')
        pBdr.append(bdr)
    pPr.append(pBdr)
    doc.add_paragraph()


def _parse_markdown_table(text: str) -> list[list[str]]:
    """
    Extract rows from a markdown table.
    Returns list of lists (first row = headers).
    Skips separator rows (---, ===).
    """
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-| :]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


# =============================================================================
# MARKDOWN BUILDER
# =============================================================================

def _build_markdown(ctx: dict) -> str:
    repo_name = os.getenv("GITHUB_REPO", "Unknown")
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sha = ctx.get("commit_sha", "")[:7]
    branch = ctx.get("branch", "")
    author = ctx.get("author", "Unknown")
    repo_type = ctx.get("repo_type", "unknown").upper()

    lines = []
    lines.append(f"# Documentation — {repo_name}")
    lines.append(f"\n> Auto-generated | Last updated: {now} | Commit: `{sha}` on `{branch}` by {author}\n")
    lines.append("---\n")

    # Repo description (bullet sections)
    repo_desc = ctx.get("repo_description", "")
    if repo_desc:
        lines.append(repo_desc)
        lines.append("\n---\n")

    # Tech stack
    tech = ctx.get("tech_stack", {})
    if tech:
        lines.append("## Tools & Tech Stack\n")
        langs = tech.get("languages", [])
        if langs:
            lines.append("**Languages:** " + ", ".join(langs) + "\n")
        libs = tech.get("frameworks_and_libs", [])
        if libs:
            lines.append("\n| Library / Framework | Category |")
            lines.append("|---|---|")
            for lib in libs:
                lines.append(f"| {lib['name']} | {lib['category']} |")
        infra = tech.get("infra", [])
        if infra:
            lines.append(f"\n**Infrastructure:** {', '.join(infra)}\n")
        lines.append(f"\n**Repository Type:** `{repo_type}`\n")
        lines.append("\n---\n")

    # Metrics
    metrics = ctx.get("metrics", [])
    if metrics:
        lines.append("## Code Quality Metrics\n")
        lines.append("| Metric | Value | Status |")
        lines.append("|---|---|---|")
        for m in metrics:
            lines.append(f"| {m.get('metric','')} | {m.get('value','')} | {m.get('status','')} |")
        lines.append("\n---\n")

    # Impact analysis table
    impact = ctx.get("impact_analysis", "")
    if impact:
        lines.append("## Impact Analysis\n")
        lines.append(impact)
        lines.append("\n---\n")

    # Combined commit table
    docs = ctx.get("documentation", "")
    if docs:
        lines.append("## Commit Change Details\n")
        lines.append(docs)
        lines.append("\n---\n")

    return "\n".join(lines)
